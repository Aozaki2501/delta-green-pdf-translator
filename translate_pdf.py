#!/usr/bin/env python3
"""
DG TRPG PDF Translator — THE MILLENNIUM (v2.0)
===============================================
Translates English PDF (dual-column TRPG layout) to Chinese Markdown/HTML/Word
using DeepSeek V4 API with TRPG-specific terminology.

v2.0 Features:
    - Context window: carries previous page context for cross-page coherence
    - Chapter detection: auto-detects headings by font size/bold for TOC generation
    - HTML output: browser-readable dual-column layout for review and print
    - Batch concurrency: translates multiple pages in parallel
    - Token cost tracking: real-time usage and cost display
    - Breakpoint resume
    - TRPG glossary support

Usage:
    python translate_pdf.py input.pdf --api-key YOUR_KEY
    python translate_pdf.py input.pdf --api-key YOUR_KEY --format html --workers 4
"""

import argparse
import hashlib
import html
import json
import os
import re
import sys
import time
import threading
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional


def configure_console_output():
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(errors="replace")
        except AttributeError:
            pass


configure_console_output()

try:
    import pymupdf  # PyMuPDF >= 1.24
except ImportError:
    try:
        import fitz as pymupdf  # PyMuPDF < 1.24
    except ImportError:
        print("Error: PyMuPDF not installed. Run: pip install pymupdf")
        sys.exit(1)

try:
    from openai import OpenAI
except ImportError:
    print("Error: openai package not installed. Run: pip install openai")
    sys.exit(1)

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


PROMPT_VERSION = "2026-05-15-preserve-layout-markers-v4"
EXTRACTOR_VERSION = "2026-05-15-card-sections-v1"
SUPPORTED_OUTPUT_FORMATS = {"markdown", "html", "word", "both", "all"}
TRANSLATION_FAILURE_PREFIX = "[Translation failed:"


def ensure_output_parent(path: str):
    parent = Path(path).expanduser().resolve().parent
    parent.mkdir(parents=True, exist_ok=True)


def normalize_page_range(start_page, end_page, total_pages: int) -> tuple[int, int]:
    try:
        start = int(start_page or 0)
    except (TypeError, ValueError) as exc:
        raise ValueError("起始页必须是整数") from exc

    try:
        end = total_pages if end_page is None else int(end_page)
    except (TypeError, ValueError) as exc:
        raise ValueError("结束页必须是整数") from exc

    if total_pages < 1:
        raise ValueError("PDF 没有可处理页面")
    if start < 0:
        raise ValueError("起始页不能小于 0")
    if start >= total_pages:
        raise ValueError(f"起始页超出范围：PDF 共 {total_pages} 页")
    if end > total_pages:
        end = total_pages
    if end <= start:
        raise ValueError("结束页必须大于起始页")
    return start, end


def is_failed_translation(text: str) -> bool:
    return bool(text and text.lstrip().startswith(TRANSLATION_FAILURE_PREFIX))


# ============================================================
# TOKEN COST TRACKER
# ============================================================

@dataclass
class TokenStats:
    """Tracks token usage and estimated cost."""
    input_tokens: int = 0
    output_tokens: int = 0
    cached_tokens: int = 0
    api_calls: int = 0
    failed_calls: int = 0
    _lock: threading.Lock = field(default_factory=threading.Lock)

    PRICE_INPUT_PER_M = 1.0
    PRICE_OUTPUT_PER_M = 4.0
    PRICE_CACHED_PER_M = 0.1

    def add(self, input_tok: int, output_tok: int, cached_tok: int = 0):
        with self._lock:
            self.input_tokens += input_tok
            self.output_tokens += output_tok
            self.cached_tokens += cached_tok
            self.api_calls += 1

    def add_failure(self):
        with self._lock:
            self.failed_calls += 1

    @property
    def total_tokens(self):
        return self.input_tokens + self.output_tokens

    @property
    def cost_yuan(self):
        cost = (
            (self.input_tokens - self.cached_tokens) * self.PRICE_INPUT_PER_M / 1_000_000 +
            self.output_tokens * self.PRICE_OUTPUT_PER_M / 1_000_000 +
            self.cached_tokens * self.PRICE_CACHED_PER_M / 1_000_000
        )
        return cost

    def summary(self) -> str:
        return (
            f"Token Stats:\n"
            f"   Input: {self.input_tokens:,} tokens\n"
            f"   Output: {self.output_tokens:,} tokens\n"
            f"   Cache hit: {self.cached_tokens:,} tokens\n"
            f"   API calls: {self.api_calls} (failed {self.failed_calls})\n"
            f"   Est. cost: Y{self.cost_yuan:.3f}"
        )


# ============================================================
# CHAPTER / HEADING DETECTION
# ============================================================

@dataclass
class HeadingInfo:
    """Stores detected heading information."""
    page_num: int
    text: str
    level: int
    y_position: float


class ChapterDetector:
    """Detects chapter/section headings by analyzing font size and weight."""

    def __init__(self):
        self.headings: list[HeadingInfo] = []
        self._font_sizes: list[float] = []

    def analyze_page(self, page_num: int, page_dict: dict):
        blocks = page_dict.get("blocks", [])
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                line_text = "".join(s["text"] for s in spans).strip()
                if not line_text or len(line_text) < 2:
                    continue
                avg_size = sum(s["size"] for s in spans) / len(spans)
                is_bold = any(s["flags"] & 2 for s in spans)
                is_all_caps = line_text == line_text.upper() and line_text != line_text.lower()
                self._font_sizes.append(avg_size)
                bbox = line.get("bbox", block["bbox"])
                h = HeadingInfo(page_num=page_num, text=line_text, level=0, y_position=bbox[1])
                h._size = avg_size
                h._bold = is_bold
                h._caps = is_all_caps
                self.headings.append(h)

    def finalize(self):
        if not self._font_sizes:
            self.headings = []
            return
        sizes = sorted(self._font_sizes)
        median_size = sizes[len(sizes) // 2]
        real_headings = []
        for h in self.headings:
            size = getattr(h, "_size", 0)
            bold = getattr(h, "_bold", False)
            caps = getattr(h, "_caps", False)
            if size >= median_size * 1.3:
                if size >= median_size * 1.8:
                    h.level = 1
                elif size >= median_size * 1.4:
                    h.level = 2
                else:
                    h.level = 3
                real_headings.append(h)
            elif bold and caps and size >= median_size * 1.1:
                h.level = 2
                real_headings.append(h)
        self.headings = real_headings

    def get_toc_markdown(self) -> str:
        if not self.headings:
            return ""
        lines = ["## Table of Contents\n"]
        for h in self.headings:
            indent = "  " * (h.level - 1)
            lines.append(f"{indent}- [{h.text}](#page-{h.page_num + 1})")
        lines.append("")
        return "\n".join(lines)

    def get_heading_for_page(self, page_num: int) -> Optional[str]:
        page_headings = [h for h in self.headings if h.page_num == page_num]
        if page_headings:
            return min(page_headings, key=lambda h: h.level).text
        return None


# ============================================================
# PDF TEXT EXTRACTION
# ============================================================

class PDFExtractor:
    """Extracts text from dual-column TRPG PDFs with intelligent layout detection."""

    def __init__(self, pdf_path: str):
        self.doc = pymupdf.open(pdf_path)
        self.total_pages = len(self.doc)
        self.chapter_detector = ChapterDetector()
        self._page_body_context: dict[int, str] = {}
        self._page_layout_notes: dict[int, list[str]] = {}

    def get_context_text(self, page_num: int) -> str:
        return self._page_body_context.get(page_num, "")

    def get_layout_notes(self, page_num: int) -> list[str]:
        return self._page_layout_notes.get(page_num, [])

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        self.close()
        return False

    def _sort_blocks_layout_aware(self, blocks, page_width, page_height=None):
        sorted_input = sorted(blocks, key=lambda b: (b["bbox"][1], b["bbox"][0]))
        non_full_blocks = [
            b for b in sorted_input
            if (
                b.get("type") == 0
                and (b["bbox"][2] - b["bbox"][0]) <= page_width * 0.6
                and not self._is_title_card_block(b, page_width, page_height)
            )
        ]
        left_count = sum(
            1 for b in non_full_blocks
            if ((b["bbox"][0] + b["bbox"][2]) / 2) < page_width / 2
        )
        right_count = len(non_full_blocks) - left_count
        if left_count < 2 or right_count < 2:
            return sorted_input

        output_blocks = []
        left_blocks = []
        right_blocks = []
        median_size = self._median_font_size(sorted_input)

        def flush_columns():
            nonlocal left_blocks, right_blocks
            output_blocks.extend(self._merge_columns_for_reading(left_blocks, right_blocks, median_size))
            left_blocks = []
            right_blocks = []

        for block in sorted_input:
            if block.get("type") != 0:
                continue

            x0, _, x1, _ = block["bbox"]
            block_width = x1 - x0
            is_full_width = block_width > page_width * 0.6
            is_title_card = self._is_title_card_block(block, page_width, page_height, median_size)

            if is_full_width or is_title_card:
                flush_columns()
                if is_title_card:
                    block = dict(block)
                    block["_dg_title_card"] = True
                output_blocks.append(block)
                continue

            block_center_x = (x0 + x1) / 2
            if block_center_x < page_width / 2:
                left_blocks.append(block)
            else:
                right_blocks.append(block)

        flush_columns()
        return output_blocks

    def _merge_columns_for_reading(self, left_blocks, right_blocks, median_size):
        left_sorted = sorted(left_blocks, key=lambda b: b["bbox"][1])
        right_sorted = sorted(right_blocks, key=lambda b: b["bbox"][1])
        if not left_sorted or not right_sorted:
            return left_sorted + right_sorted

        merged = []
        right_idx = 0

        for idx, left_block in enumerate(left_sorted):
            merged.append(left_block)
            next_left = left_sorted[idx + 1] if idx + 1 < len(left_sorted) else None
            if not next_left:
                continue

            if not self._is_heading_block(next_left, median_size):
                continue
            if self._ends_like_complete_sentence(self._extract_block_text(left_block)):
                continue

            heading_y = next_left["bbox"][1]
            while right_idx < len(right_sorted):
                right_block = right_sorted[right_idx]
                if right_block["bbox"][1] >= heading_y:
                    break
                right_text = self._extract_block_text(right_block)
                if not self._starts_with_lowercase(right_text):
                    break
                merged.append(right_block)
                right_idx += 1

        merged.extend(right_sorted[right_idx:])
        return merged

    def _median_font_size(self, blocks):
        sizes = []
        for block in blocks:
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    size = span.get("size")
                    if size:
                        sizes.append(size)
        if not sizes:
            return 10
        sizes.sort()
        return sizes[len(sizes) // 2]

    def _block_avg_font_size(self, block):
        sizes = [
            span.get("size", 0)
            for line in block.get("lines", [])
            for span in line.get("spans", [])
            if span.get("size")
        ]
        if not sizes:
            return 0
        return sum(sizes) / len(sizes)

    def _block_width(self, block):
        x0, _, x1, _ = block["bbox"]
        return x1 - x0

    def _block_height(self, block):
        _, y0, _, y1 = block["bbox"]
        return y1 - y0

    def _block_center_x(self, block):
        x0, _, x1, _ = block["bbox"]
        return (x0 + x1) / 2

    def _block_line_count(self, block):
        return sum(
            1 for line in block.get("lines", [])
            if self._extract_line_text(line).strip()
        )

    def _rect_from_bbox(self, bbox):
        return pymupdf.Rect(*bbox)

    def _rect_contains_block(self, rect, block, tolerance=4):
        x0, y0, x1, y1 = block["bbox"]
        return (
            x0 >= rect.x0 - tolerance
            and y0 >= rect.y0 - tolerance
            and x1 <= rect.x1 + tolerance
            and y1 <= rect.y1 + tolerance
        )

    def _rects_touch_or_overlap(self, left, right, tolerance=10):
        return not (
            left.x1 < right.x0 - tolerance
            or right.x1 < left.x0 - tolerance
            or left.y1 < right.y0 - tolerance
            or right.y1 < left.y0 - tolerance
        )

    def _union_rect(self, rects):
        x0 = min(rect.x0 for rect in rects)
        y0 = min(rect.y0 for rect in rects)
        x1 = max(rect.x1 for rect in rects)
        y1 = max(rect.y1 for rect in rects)
        return pymupdf.Rect(x0, y0, x1, y1)

    def _block_fonts(self, block):
        return {
            span.get("font", "")
            for line in block.get("lines", [])
            for span in line.get("spans", [])
            if span.get("font")
        }

    def _is_monospace_block(self, block):
        fonts = self._block_fonts(block)
        return any("VT323" in font or "Mono" in font or "Courier" in font for font in fonts)

    def _line_words(self, line):
        words = []
        for span in line.get("spans", []):
            text = span.get("text", "").strip()
            if not text:
                continue
            x0, y0, x1, y1 = span.get("bbox", line.get("bbox", (0, 0, 0, 0)))
            words.append({"text": text, "x": x0, "y": y0, "bbox": (x0, y0, x1, y1)})
        return words

    def _extract_monospace_lines(self, block):
        lines = []
        for line in block.get("lines", []):
            words = self._line_words(line)
            if not words:
                continue
            words.sort(key=lambda word: word["x"])
            lines.append(words)
        return lines

    def _block_to_markdown_table(self, block):
        return self._blocks_to_markdown_table([block])

    def _blocks_to_markdown_table(self, blocks):
        row_candidates = []
        for block in blocks:
            block_words = []
            for words in self._extract_monospace_lines(block):
                row_text = " ".join(word["text"] for word in words)
                if re.fullmatch(r"[_+\-\s|]+", row_text):
                    continue
                block_words.extend(words)
            if block_words:
                row_candidates.append(sorted(block_words, key=lambda word: (word["y"], word["x"])))

        if len(row_candidates) < 2:
            return None

        header_words = row_candidates[0]
        header = [re.sub(r"^[| ]+|[| ]+$", "", word["text"]).strip() for word in header_words]
        header = [cell for cell in header if cell]
        if len(header) < 2:
            return None

        col_count = len(header)
        col_positions = sorted(word["x"] for word in header_words[:col_count])
        table_rows = []
        for words in row_candidates[1:]:
            cells = ["" for _ in range(col_count)]
            for word in words:
                idx = min(range(col_count), key=lambda i: abs(word["x"] - col_positions[i]))
                clean_word = re.sub(r"^[| ]+|[| ]+$", "", word["text"]).strip()
                if clean_word:
                    cells[idx] = (cells[idx] + " " + clean_word).strip()
            if any(cells):
                table_rows.append(cells)

        if not table_rows:
            return None

        markdown = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join("---" for _ in header) + " |",
        ]
        for cells in table_rows:
            markdown.append("| " + " | ".join(cells) + " |")
        return "\n".join(markdown)

    def _is_table_block(self, block, page_width):
        if not self._is_monospace_block(block):
            return False
        x0, _, x1, _ = block["bbox"]
        if x1 - x0 < page_width * 0.55:
            return False
        text = self._extract_block_text(block)
        return "|" in text or "CLUE" in text.upper() or re.search(r"_{8,}|\+[-_ ]+", text)

    def _is_contents_block(self, block):
        if not self._is_monospace_block(block):
            return False
        text = self._extract_block_text(block)
        return bool(re.search(r"\bContents\b", text, re.IGNORECASE)) or (
            text.count(".") >= 20 and bool(re.search(r"\.{4,}\s*\d{1,3}", text))
        )

    def _is_handout_block(self, block):
        if not self._is_monospace_block(block):
            return False
        text = self._extract_block_text(block)
        if not text:
            return False
        return bool(re.search(r"\b(SUBJECT|Records?|Stories?|Profile)\b", text, re.IGNORECASE))

    def _has_card_label(self, text: str) -> bool:
        patterns = [
            r"\b(?:YELLOW|GREEN|RED|BLUE|WHITE|BLACK)\s+CARD\b",
            r"\bPLAYER\s+AID\b",
            r"\bSUBJECT\s*:",
            r"\bPROFILE\s+OF\b",
            r"\b(?:Birth|Medical|Police|USMC|Military|News|School|Juvenile)\s+Records?\b",
            r"^\s*(?:Timeline|Briefing|Report|Memo|Evidence|Clue|Handout|Photograph|Letter|Note)\b",
        ]
        return any(re.search(pattern, text, re.IGNORECASE) for pattern in patterns)

    def _is_card_text_block(self, block, page_width, page_height, median_size=None):
        text = self._extract_block_text(block).strip()
        if not text:
            return False
        if self._is_contents_block(block) or self._is_table_block(block, page_width):
            return False
        width = self._block_width(block)
        line_count = self._block_line_count(block)
        avg_size = self._block_avg_font_size(block)
        median_size = median_size or avg_size or 10
        if self._is_handout_block(block):
            return True
        if self._has_card_label(text) and (width >= page_width * 0.28 or line_count >= 2):
            return True
        if (
            self._is_monospace_block(block)
            and line_count >= 4
            and width >= page_width * 0.42
            and self._block_height(block) >= page_height * 0.08
        ):
            return True
        if (
            width >= page_width * 0.68
            and line_count >= 5
            and avg_size <= median_size * 1.35
            and self._has_card_label(text[:300])
        ):
            return True
        return False

    def _visual_card_regions(self, page, content_blocks, page_width, page_height):
        page_area = page_width * page_height
        regions = []

        for drawing in page.get_drawings():
            rect = drawing.get("rect")
            if not rect:
                continue
            area = rect.width * rect.height
            if area < page_area * 0.025 or area > page_area * 0.78:
                continue
            if rect.width < page_width * 0.24 or rect.height < page_height * 0.08:
                continue
            regions.append(rect)

        for image in page.get_images(full=True):
            xref = image[0]
            for rect in page.get_image_rects(xref):
                area = rect.width * rect.height
                if area < page_area * 0.035 or area > page_area * 0.78:
                    continue
                if rect.width < page_width * 0.28 or rect.height < page_height * 0.10:
                    continue
                regions.append(rect)

        median_size = self._median_font_size(content_blocks)
        accepted = []
        for rect in regions:
            inside = [
                block for block in content_blocks
                if self._rect_contains_block(rect, block, tolerance=8)
            ]
            if not inside:
                continue
            has_card_text = any(
                self._is_card_text_block(block, page_width, page_height, median_size)
                for block in inside
            )
            monospace_lines = sum(
                self._block_line_count(block) for block in inside
                if self._is_monospace_block(block)
            )
            if has_card_text or monospace_lines >= 4:
                accepted.append(rect)

        merged = []
        for rect in sorted(accepted, key=lambda item: (item.y0, item.x0)):
            for idx, existing in enumerate(merged):
                if self._rects_touch_or_overlap(existing, rect, tolerance=14):
                    merged[idx] = self._union_rect([existing, rect])
                    break
            else:
                merged.append(rect)
        return merged

    def _group_card_blocks(self, card_blocks, page_width, page_height):
        groups = []
        for block in sorted(card_blocks, key=lambda item: (item["bbox"][1], item["bbox"][0])):
            rect = self._rect_from_bbox(block["bbox"])
            placed = False
            for group in groups:
                group_rect = self._union_rect([self._rect_from_bbox(b["bbox"]) for b in group])
                same_band = abs(rect.y0 - group_rect.y1) <= page_height * 0.08
                horizontal_overlap = min(rect.x1, group_rect.x1) - max(rect.x0, group_rect.x0)
                overlap_ratio = horizontal_overlap / max(min(rect.width, group_rect.width), 1)
                if same_band and overlap_ratio >= 0.35:
                    group.append(block)
                    placed = True
                    break
            if not placed:
                groups.append([block])
        return groups

    def _split_card_blocks(self, page, content_blocks, page_width, page_height):
        median_size = self._median_font_size(content_blocks)
        card_ids = set()
        card_groups = []
        notes = []

        for rect in self._visual_card_regions(page, content_blocks, page_width, page_height):
            group = [
                block for block in content_blocks
                if id(block) not in card_ids and self._rect_contains_block(rect, block, tolerance=8)
            ]
            if not group:
                continue
            for block in group:
                card_ids.add(id(block))
            card_groups.append(group)

        loose_card_blocks = [
            block for block in content_blocks
            if id(block) not in card_ids
            and self._is_card_text_block(block, page_width, page_height, median_size)
        ]
        for group in self._group_card_blocks(loose_card_blocks, page_width, page_height):
            for block in group:
                card_ids.add(id(block))
            card_groups.append(group)

        body_blocks = [
            block for block in content_blocks
            if id(block) not in card_ids
        ]
        if card_groups:
            card_blocks_count = sum(len(group) for group in card_groups)
            notes.append(f"{len(card_groups)} card section(s), {card_blocks_count} block(s)")
        return body_blocks, card_groups, notes

    def _is_heading_block(self, block, median_size):
        text = self._extract_block_text(block).strip()
        if not text or len(text) > 90:
            return False
        if re.search(r"[.!?。！？]$", text):
            return False
        return self._block_avg_font_size(block) >= median_size * 1.25

    def _is_title_card_block(self, block, page_width, page_height=None, median_size=None):
        text = self._extract_block_text(block).strip()
        text = re.sub(r"^#\s*", "", text).strip()
        if not text or len(text) > 120:
            return False
        if re.search(r"[.!?。！？]$", text):
            return False

        x0, y0, x1, _ = block["bbox"]
        center_x = (x0 + x1) / 2
        page_center = page_width / 2
        avg_size = self._block_avg_font_size(block)
        size_floor = (median_size * 2.0) if median_size else 22
        near_top = True if page_height is None else y0 < page_height * 0.22
        centered = abs(center_x - page_center) <= page_width * 0.22
        return near_top and centered and avg_size >= size_floor

    def _ends_like_complete_sentence(self, text):
        return bool(re.search(r"[.!?。！？”\"’')\]]\s*$", text.strip()))

    def _starts_with_lowercase(self, text):
        match = re.search(r"[A-Za-z]", text.strip())
        return bool(match and match.group(0).islower())

    def _extract_line_text(self, line):
        spans_text = []
        last_span_norm = ""
        for span in line.get("spans", []):
            span_text = span["text"]
            span_norm = re.sub(r"\s+", " ", span_text).strip().lower()
            if span_norm and span_norm == last_span_norm:
                continue
            spans_text.append(span_text)
            last_span_norm = span_norm
        return "".join(spans_text).strip()

    def _line_avg_font_size(self, line):
        sizes = [
            span.get("size", 0)
            for span in line.get("spans", [])
            if span.get("size")
        ]
        if not sizes:
            return 0
        return sum(sizes) / len(sizes)

    def _join_line_into_paragraph(self, paragraph, line_text):
        if not paragraph:
            return line_text
        tail = paragraph.rstrip()
        if tail.endswith("-") and re.search(r"[A-Za-z]-$", tail) and re.match(r"^[A-Za-z]", line_text):
            return tail[:-1] + line_text
        return tail + " " + line_text

    def _extract_block_text(self, block):
        raw_lines = []
        for line in block.get("lines", []):
            line_text = self._extract_line_text(line)
            if not line_text:
                continue
            if raw_lines and line_text == raw_lines[-1]["text"]:
                continue
            raw_lines.append({
                "text": line_text,
                "bbox": line.get("bbox", block["bbox"]),
                "size": self._line_avg_font_size(line),
            })
        if not raw_lines:
            return ""

        body_sizes = sorted(l["size"] for l in raw_lines if l["size"])
        median_size = body_sizes[len(body_sizes) // 2] if body_sizes else 10
        left_edge = min(l["bbox"][0] for l in raw_lines)

        paragraphs = []
        current = ""

        def flush():
            nonlocal current
            if current.strip():
                paragraphs.append(current.strip())
                current = ""

        for idx, line in enumerate(raw_lines):
            text = line["text"]
            indent = line["bbox"][0] - left_edge
            visible = re.sub(r"\s+", "", text)
            is_heading = (
                line["size"] >= median_size * 1.6
                and 2 <= len(visible) <= 40
                and not re.search(r"[.!?。！？]$", text)
            )
            is_indented_para = idx > 0 and indent >= max(10, median_size * 1.2)

            if is_heading:
                flush()
                paragraphs.append(text)
                continue
            if is_indented_para:
                flush()

            current = self._join_line_into_paragraph(current, text)

        flush()
        text = "\n\n".join(paragraphs)
        if block.get("_dg_title_card"):
            title = re.sub(r"\s+", " ", text).strip()
            return f"# {title}" if title else ""
        return text

    def _is_header_footer(self, block, page_height, margin_ratio=0.08):
        top_margin = page_height * margin_ratio
        bottom_margin = page_height * (1 - margin_ratio)
        block_y = block["bbox"][1]
        block_y_bottom = block["bbox"][3]
        text = self._extract_block_text(block).strip()
        if not text:
            return True

        compact = re.sub(r"\s+", " ", text)
        normalized = re.sub(r"[^A-Z0-9 ]+", "", compact.upper()).strip()
        in_top = block_y_bottom < top_margin
        in_bottom = block_y > bottom_margin
        in_margin = in_top or in_bottom

        if re.fullmatch(r"\d{1,4}", compact):
            return True

        running_titles = ("DELTA GREEN", "PISCES", "THE MILLENNIUM", "THE NEW AGE", "THE LABYRINTH")
        if in_margin and any(title in normalized for title in running_titles):
            return True
        if in_margin and "//" in compact:
            return True

        if in_bottom and len(compact) <= 80:
            return True
        return False

    def detect_page_layout(self, page_num: int) -> str:
        """Return 'handout', 'single', or 'columns' for source page layout."""
        page = self.doc[page_num]
        page_width = page.rect.width
        page_height = page.rect.height
        page_dict = page.get_text("dict", flags=pymupdf.TEXT_PRESERVE_WHITESPACE)
        content_blocks = [
            b for b in page_dict.get("blocks", [])
            if b.get("type") == 0 and not self._is_header_footer(b, page_height)
        ]
        if not content_blocks:
            return "columns"

        if any(self._is_contents_block(block) for block in content_blocks):
            return "toc"

        handout_blocks = [
            block for block in content_blocks
            if self._is_handout_block(block)
        ]
        top_blocks = sorted(content_blocks, key=lambda block: (block["bbox"][1], block["bbox"][0]))
        top_text = self._extract_block_text(top_blocks[0]) if top_blocks else ""
        if len(handout_blocks) >= 3 or re.match(r"\s*Player Aid\b", top_text, re.IGNORECASE):
            return "handout"

        left_count = 0
        right_count = 0
        full_width_height = 0
        total_height = 0

        for block in content_blocks:
            x0, y0, x1, y1 = block["bbox"]
            width = x1 - x0
            height = max(0, y1 - y0)
            center = (x0 + x1) / 2
            total_height += height

            spans_most_page = (
                width >= page_width * 0.72
                and x0 <= page_width * 0.20
                and x1 >= page_width * 0.80
            )
            if spans_most_page:
                full_width_height += height
                continue

            if width <= page_width * 0.62:
                if center < page_width / 2:
                    left_count += 1
                else:
                    right_count += 1

        has_two_column_signal = left_count >= 1 and right_count >= 1
        full_width_ratio = full_width_height / max(total_height, 1)
        if full_width_ratio >= 0.45:
            return "single"
        if has_two_column_signal:
            return "columns"
        if len(content_blocks) <= 3 and full_width_height > page_height * 0.18:
            return "single"
        return "columns"

    def _clean_text(self, text):
        text = re.sub(r"(\w)-\s*\n\s*(\w)", r"\1\2", text)
        text = re.sub(r"  +", " ", text)
        text = re.sub(r"^\s*\d{1,3}\s*$", "", text, flags=re.MULTILINE)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _layout_notes_for_page(self, layout: str, content_blocks, page_width: float) -> list[str]:
        notes = [f"layout: {layout}"]
        table_count = sum(1 for block in content_blocks if self._is_table_block(block, page_width))
        handout_count = sum(1 for block in content_blocks if self._is_handout_block(block))
        if any(self._is_contents_block(block) for block in content_blocks):
            notes.append("contents page preserved as TOC")
        if table_count:
            notes.append(f"{table_count} table-like block(s)")
        if handout_count:
            notes.append(f"{handout_count} handout block(s)")
        return notes

    def _context_from_extracted_text(self, text: str, layout: str) -> str:
        if layout == "toc":
            return ""
        context_lines = []
        in_fenced_block = False
        for raw_line in text.splitlines():
            line = raw_line.strip()
            if line.startswith("```"):
                in_fenced_block = not in_fenced_block
                continue
            if in_fenced_block:
                continue
            if line == "[[TOC]]":
                continue
            if line.startswith("|") and line.count("|") >= 2:
                continue
            if line.startswith(">"):
                line = line.lstrip(">").strip()
            context_lines.append(line)
        return self._clean_text("\n".join(context_lines))[-1200:]

    def _quote_card_text(self, text: str) -> str:
        quoted = []
        for line in text.splitlines():
            quoted.append("> " + line.strip() if line.strip() else ">")
        return "\n".join(quoted).strip()

    def _blocks_to_extracted_text(self, blocks, page_width, page_height, layout_aware=True,
                                  mark_handouts=True) -> str:
        if not blocks:
            return ""
        if layout_aware:
            sorted_blocks = self._sort_blocks_layout_aware(blocks, page_width, page_height)
        else:
            sorted_blocks = sorted(blocks, key=lambda item: (item["bbox"][1], item["bbox"][0]))

        processed_blocks = []
        idx = 0
        while idx < len(sorted_blocks):
            block = sorted_blocks[idx]
            if self._is_table_block(block, page_width):
                table_blocks = []
                while idx < len(sorted_blocks) and self._is_monospace_block(sorted_blocks[idx]):
                    table_blocks.append(sorted_blocks[idx])
                    idx += 1
                table_text = self._blocks_to_markdown_table(table_blocks)
                if not table_text:
                    table_text = "\n".join(
                        self._extract_block_text(b).strip()
                        for b in table_blocks
                        if self._extract_block_text(b).strip()
                    )
                processed_blocks.append({"text": table_text, "title_card": False})
                continue

            if mark_handouts and self._is_handout_block(block):
                text = self._extract_block_text(block).strip()
                if text:
                    text = self._quote_card_text(text)
                processed_blocks.append({"text": text, "title_card": False})
                idx += 1
                continue

            processed_blocks.append({
                "text": self._extract_block_text(block).strip(),
                "title_card": bool(block.get("_dg_title_card")),
            })
            idx += 1

        paragraphs = []
        current_para = ""
        current_is_title_card = False

        for item in processed_blocks:
            text = item["text"].strip()
            if not text:
                continue
            is_title_card = item["title_card"]

            if not current_para:
                current_para = text
                current_is_title_card = is_title_card
                continue

            if current_is_title_card and is_title_card:
                left = re.sub(r"^#\s*", "", current_para).strip()
                right = re.sub(r"^#\s*", "", text).strip()
                current_para = "# " + " ".join(part for part in (left, right) if part)
                continue

            current_tail = current_para.rstrip()
            first_alpha = re.search(r"[A-Za-z]", text)
            starts_lower = bool(first_alpha and first_alpha.group(0).islower())
            joins_from_punctuation = current_tail.endswith((",", ":", ";", "-"))

            if joins_from_punctuation or starts_lower:
                if current_tail.endswith("-"):
                    current_para = current_tail[:-1].rstrip() + text
                else:
                    current_para = current_tail + " " + text
                current_is_title_card = current_is_title_card and is_title_card
            else:
                paragraphs.append(current_para)
                current_para = text
                current_is_title_card = is_title_card

        if current_para:
            paragraphs.append(current_para)
        return self._clean_text("\n\n".join(paragraphs))

    def extract_page(self, page_num: int) -> str:
        page = self.doc[page_num]
        page_width = page.rect.width
        page_height = page.rect.height
        page_dict = page.get_text("dict", flags=pymupdf.TEXT_PRESERVE_WHITESPACE)
        blocks = page_dict.get("blocks", [])
        self.chapter_detector.analyze_page(page_num, page_dict)
        if not blocks:
            self._page_body_context[page_num] = ""
            self._page_layout_notes[page_num] = ["empty page"]
            return ""
        content_blocks = [
            b for b in blocks
            if b.get("type") == 0 and not self._is_header_footer(b, page_height)
        ]
        if not content_blocks:
            self._page_body_context[page_num] = ""
            self._page_layout_notes[page_num] = ["no content blocks"]
            return ""
        layout = self.detect_page_layout(page_num)
        self._page_layout_notes[page_num] = self._layout_notes_for_page(layout, content_blocks, page_width)
        if layout == "toc":
            toc_text = self._extract_contents_page(content_blocks)
            clean_toc = self._clean_text(toc_text)
            self._page_body_context[page_num] = ""
            return clean_toc

        body_blocks, card_groups, card_notes = self._split_card_blocks(
            page, content_blocks, page_width, page_height
        )
        self._page_layout_notes[page_num].extend(card_notes)

        body_text = self._blocks_to_extracted_text(
            body_blocks, page_width, page_height, layout_aware=True, mark_handouts=True
        )
        sections = []
        if body_text:
            sections.append(body_text)
        for group in sorted(card_groups, key=lambda group: min(block["bbox"][1] for block in group)):
            card_text = self._blocks_to_extracted_text(
                group, page_width, page_height, layout_aware=False, mark_handouts=False
            )
            if card_text:
                sections.append(self._quote_card_text(card_text))

        clean_text = self._clean_text("\n\n".join(sections))
        self._page_body_context[page_num] = self._context_from_extracted_text(body_text, layout)
        return clean_text

    def _extract_contents_page(self, content_blocks):
        toc_blocks = [
            block for block in content_blocks
            if self._is_monospace_block(block)
        ]
        title_blocks = [block for block in toc_blocks if self._is_contents_block(block) and "Contents" in self._extract_block_text(block)]
        body_blocks = [block for block in toc_blocks if block not in title_blocks]
        body_blocks = sorted(body_blocks, key=lambda b: (b["bbox"][0], b["bbox"][1]))

        parts = []
        if title_blocks:
            title = self._extract_block_text(sorted(title_blocks, key=lambda b: b["bbox"][1])[0])
            title = re.sub(r"_+", "", title).strip()
            if title:
                parts.append(f"[[TOC]]\n# {title}")
        else:
            parts.append("[[TOC]]")

        for block in body_blocks:
            text = self._extract_contents_block_lines(block).strip()
            if text:
                parts.append("```toc\n" + text + "\n```")
        return "\n\n".join(parts)

    def _extract_contents_block_lines(self, block):
        lines = []
        for line in block.get("lines", []):
            text = self._extract_line_text(line)
            text = re.sub(r"\s+", " ", text).rstrip()
            if text:
                lines.append(text)
        return "\n".join(lines)

    def finalize_chapters(self):
        self.chapter_detector.finalize()

    def close(self):
        if self.doc is not None:
            self.doc.close()
            self.doc = None


# ============================================================
# GLOSSARY LOADER
# ============================================================

def load_glossary(glossary_path: str) -> dict:
    glossary = {}
    if not glossary_path or not os.path.exists(glossary_path):
        return glossary
    with open(glossary_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            if "\t" in line:
                parts = line.split("\t", 1)
            else:
                parts = re.split(r"\s{2,}", line, maxsplit=1)
            if len(parts) == 2:
                chinese = parts[0].strip()
                english = parts[1].strip()
                if english and chinese:
                    glossary[english] = chinese
    return glossary


def file_sha256(path: str) -> str:
    if not path or not os.path.exists(path):
        return ""
    digest = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def build_progress_metadata(pdf_path: str, glossary_path: Optional[str], model: str,
                            start_page: int, end_page: int | None) -> dict:
    return {
        "schema": 1,
        "pdf_sha256": file_sha256(pdf_path),
        "glossary_sha256": file_sha256(glossary_path) if glossary_path else "",
        "model": model,
        "prompt_version": PROMPT_VERSION,
        "extractor_version": EXTRACTOR_VERSION,
        "start_page": start_page,
        "end_page": end_page,
    }


def compare_progress_metadata(expected: dict, actual: dict) -> list[str]:
    if not expected:
        return []
    if not actual:
        return ["进度文件缺少版本指纹"]

    mismatches = []
    for key, expected_value in expected.items():
        actual_value = actual.get(key)
        if actual_value != expected_value:
            mismatches.append(f"{key}: {actual_value!r} -> {expected_value!r}")
    return mismatches


def parse_page_selection(selection: str, total_pages: int) -> set[int]:
    """Parse 1-based page specs such as '8, 12-15' into zero-based page indexes."""
    pages = set()
    if not selection or not selection.strip():
        return pages

    for raw_part in re.split(r"[,\s，、]+", selection.strip()):
        part = raw_part.strip()
        if not part:
            continue
        try:
            if "-" in part:
                start_text, end_text = part.split("-", 1)
                if not start_text.strip() or not end_text.strip():
                    raise ValueError
                start = int(start_text)
                end = int(end_text)
                if start > end:
                    start, end = end, start
                page_numbers = range(start, end + 1)
            else:
                page_numbers = [int(part)]
        except ValueError as exc:
            raise ValueError(f"无法解析页码片段：{part!r}") from exc

        for page_number in page_numbers:
            if page_number < 1 or page_number > total_pages:
                raise ValueError(f"页码 {page_number} 超出范围 1-{total_pages}")
            pages.add(page_number - 1)
    return pages


def find_relevant_glossary_terms(text: str, glossary: dict) -> dict:
    matches = []
    for eng, chn in sorted(glossary.items(), key=lambda item: len(item[0]), reverse=True):
        pattern = re.compile(
            r"(?<![A-Za-z0-9])" + re.escape(eng) + r"(?![A-Za-z0-9])",
            re.IGNORECASE,
        )
        for match in pattern.finditer(text):
            matches.append((match.start(), match.end(), eng, chn))

    selected = []
    occupied_spans = []
    for start, end, eng, chn in matches:
        if any(start < occupied_end and end > occupied_start for occupied_start, occupied_end in occupied_spans):
            continue
        selected.append((eng, chn))
        occupied_spans.append((start, end))

    relevant = {}
    for eng, chn in selected:
        relevant[eng] = chn
    return relevant


def _find_unlisted_proper_nouns(text: str, glossary_hits: dict) -> list[str]:
    known = {term.lower() for term in glossary_hits}
    stopwords = {
        "A", "An", "And", "Are", "As", "At", "Be", "But", "By", "For", "From", "He",
        "Her", "His", "If", "In", "Into", "Is", "It", "Its", "Of", "On", "Or", "She",
        "The", "Their", "They", "This", "To", "Was", "Were", "When", "With", "You",
        "Chapter", "Page", "Table", "Figure",
    }
    candidates = {}
    pattern = re.compile(r"\b(?:[A-Z][A-Za-z'’.-]+)(?:\s+(?:of|the|and|&|[A-Z][A-Za-z'’.-]+))*\b")
    for match in pattern.finditer(text):
        candidate = match.group(0).strip(" -.,:;!?()[]{}\"“”")
        if len(candidate) < 3 or candidate in stopwords:
            continue
        if candidate.isupper() and len(candidate) <= 6:
            continue
        if candidate.lower() in known:
            continue
        candidates[candidate] = candidates.get(candidate, 0) + 1
    return [
        term for term, _ in sorted(candidates.items(), key=lambda item: (-item[1], item[0].lower()))[:20]
    ]


def build_glossary_report(pages_text: dict, glossary: dict, title: str = "") -> str:
    lines = [
        f"# {title} — 术语命中报告" if title else "# 术语命中报告",
        "",
        "本报告基于提取后的英文原文生成，用于检查每页实际命中的术语。",
        "",
    ]
    if not glossary:
        lines.append("未使用术语表。")
        return "\n".join(lines)

    summary = {}
    page_reports = []
    missing_candidates = {}

    for page_num in sorted(pages_text):
        text = pages_text.get(page_num, "")
        hits = find_relevant_glossary_terms(text, glossary)
        for eng, chn in hits.items():
            summary.setdefault(eng, {"chinese": chn, "pages": set()})
            summary[eng]["pages"].add(page_num + 1)
        missing = _find_unlisted_proper_nouns(text, hits)
        for term in missing:
            missing_candidates.setdefault(term, set()).add(page_num + 1)
        page_reports.append((page_num + 1, hits, missing))

    lines.append("## 汇总")
    lines.append("")
    if summary:
        for eng, info in sorted(summary.items(), key=lambda item: item[0].lower()):
            pages = _format_page_ranges([p - 1 for p in info["pages"]])
            lines.append(f"- `{eng}` -> `{info['chinese']}`；页：{pages}")
    else:
        lines.append("- 未命中任何术语。")

    lines.append("")
    lines.append("## 逐页命中")
    lines.append("")
    for page_num, hits, missing in page_reports:
        lines.append(f"### 第 {page_num} 页")
        if hits:
            for eng, chn in sorted(hits.items(), key=lambda item: item[0].lower()):
                lines.append(f"- `{eng}` -> `{chn}`")
        else:
            lines.append("- 无术语命中")
        if missing:
            lines.append(f"- 疑似未收录专名：{', '.join(missing[:10])}")
        lines.append("")

    lines.append("## 疑似未收录专名")
    lines.append("")
    if missing_candidates:
        for term, pages in sorted(missing_candidates.items(), key=lambda item: item[0].lower())[:100]:
            page_text = _format_page_ranges([p - 1 for p in pages])
            lines.append(f"- `{term}`；页：{page_text}")
    else:
        lines.append("- 暂无。")
    lines.append("")
    return "\n".join(lines)


def write_glossary_report(pages_text: dict, glossary: dict, report_output: str, title: str = ""):
    ensure_output_parent(report_output)
    report = build_glossary_report(pages_text, glossary, title)
    with open(report_output, "w", encoding="utf-8") as f:
        f.write(report)


# ============================================================
# TRANSLATOR — DeepSeek V4 API with Context Window
# ============================================================

class Translator:
    """Translates text using DeepSeek V4 API with context window and cost tracking."""

    SYSTEM_PROMPT = """You are a professional TRPG translator working on Delta Green sourcebooks.

Translation rules:
0. If the source starts with [[TOC]], preserve the table of contents structure. Do not translate entries; keep dotted leaders and page numbers.
1. Follow the glossary strictly for proper nouns. If glossary entries overlap, the longest matching phrase wins.
2. Keep untranslated: dice notations (1D6, 3D6), attributes (STR, CON, DEX, INT, POW, CHA, SAN, WP, HP), skill checks (1/1D6 SAN), abbreviations (FBI, CIA, MJ-12, A-Cell).
3. Output in Markdown format with ## headings, - bullet lists, paragraph spacing.
4. Professional, fluent Chinese. Maintain horror atmosphere. Precise rule descriptions.
5. Keep the Chinese translation concise. Do not expand, explain, embellish, or add information not present in the source.
6. Do not translate page headers, footers, page numbers, or running titles such as DELTA GREEN, PISCES, and THE MILLENNIUM when they appear as standalone navigation text.
7. Preserve section and subsection headings. Output article titles as ## headings and smaller section headings as ### headings.
8. If decorative title-card text, drop shadows, or stylized text is extracted twice, translate it only once.
9. If OCR errors/garbled text exists, infer meaning from context. Mark unreadable as [damaged].
10. If previous context is provided, ensure continuity. Do not re-translate previous content.
11. Preserve Markdown tables as Markdown tables. Translate cell text but keep the same columns.
12. Preserve blockquotes starting with > for handouts/player aids/cards. Do not merge blockquoted card text into the surrounding body text.

{glossary_section}"""

    def __init__(self, api_key: str, model: str = "deepseek-v4-pro",
                 base_url: str = "https://api.deepseek.com", stats: TokenStats = None):
        if not api_key or not str(api_key).strip():
            raise ValueError("API Key 不能为空")
        if not model or not str(model).strip():
            raise ValueError("模型名称不能为空")
        self.client = OpenAI(api_key=api_key, base_url=base_url)
        self.model = model
        self.glossary = {}
        self.retry_count = 3
        self.retry_delay = 5
        self.stats = stats or TokenStats()

    def set_glossary(self, glossary: dict):
        self.glossary = glossary

    def _build_glossary_for_chunk(self, text: str) -> str:
        if not self.glossary:
            return ""
        relevant = self._find_relevant_glossary_terms(text)
        if not relevant:
            return ""
        glossary_lines = [f"   - {eng} -> {chn}" for eng, chn in relevant.items()]
        return "\nGlossary (this section):\n" + "\n".join(glossary_lines)

    def _find_relevant_glossary_terms(self, text: str) -> dict:
        return find_relevant_glossary_terms(text, self.glossary)

    def translate_chunk(self, text: str, page_num: int = None, prev_context: str = "") -> str:
        if not text.strip():
            return ""
        glossary_section = self._build_glossary_for_chunk(text)
        system_prompt = self.SYSTEM_PROMPT.format(glossary_section=glossary_section)

        page_info = f" (page {page_num + 1})" if page_num is not None else ""
        if prev_context:
            user_prompt = (
                f"[Previous context - DO NOT translate, for reference only]\n"
                f"{prev_context}\n\n---\n\n"
                f"Translate the following{page_info}:\n\n{text}"
            )
        else:
            user_prompt = f"Translate the following{page_info}:\n\n{text}"

        for attempt in range(self.retry_count):
            try:
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=0.3,
                    max_tokens=4096,
                )
                usage = response.usage
                if usage:
                    cached = getattr(usage, "prompt_cache_hit_tokens", 0) or 0
                    self.stats.add(
                        getattr(usage, "prompt_tokens", 0) or 0,
                        getattr(usage, "completion_tokens", 0) or 0,
                        cached,
                    )
                if not response.choices:
                    raise RuntimeError("API 返回空 choices")
                content = response.choices[0].message.content or ""
                content = content.strip()
                if not content:
                    raise RuntimeError("API 返回空译文")
                return content
            except Exception as e:
                self.stats.add_failure()
                if attempt < self.retry_count - 1:
                    wait = self.retry_delay * (attempt + 1)
                    print(f"\n  API error (attempt {attempt+1}/{self.retry_count}): {e}")
                    print(f"     Retrying in {wait}s...", end="", flush=True)
                    time.sleep(wait)
                else:
                    print(f"\n  API failed permanently: {e}")
                    return f"{TRANSLATION_FAILURE_PREFIX} {e}]\n\nOriginal:\n{text[:200]}..."
        return ""


# ============================================================
# PROGRESS TRACKER
# ============================================================

class ProgressTracker:
    def __init__(self, progress_file: str, expected_metadata: Optional[dict] = None,
                 reuse_mismatched: bool = False):
        self.progress_file = progress_file
        self.expected_metadata = expected_metadata or {}
        self.reuse_mismatched = reuse_mismatched
        self.metadata = dict(self.expected_metadata)
        self.metadata_mismatches: list[str] = []
        self.ignored_existing_progress = False
        self.completed_pages = set()
        self.translations = {}
        self._lock = threading.Lock()
        self._load()

    def _load(self):
        if os.path.exists(self.progress_file):
            try:
                with open(self.progress_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if not isinstance(data, dict):
                    raise ValueError("progress root must be an object")
                loaded_metadata = data.get("metadata", {})
                self.metadata_mismatches = compare_progress_metadata(
                    self.expected_metadata,
                    loaded_metadata,
                )
                if self.metadata_mismatches and not self.reuse_mismatched:
                    self.ignored_existing_progress = True
                    self.completed_pages = set()
                    self.translations = {}
                    self.metadata = dict(self.expected_metadata)
                    print("Progress metadata mismatch, ignoring cached translations")
                    return

                self.metadata = loaded_metadata or dict(self.expected_metadata)
                self.completed_pages = {
                    int(page_num)
                    for page_num in data.get("completed_pages", [])
                    if str(page_num).isdigit()
                }
                loaded_translations = data.get("translations", {})
                self.translations = (
                    loaded_translations if isinstance(loaded_translations, dict) else {}
                )
                print(f"Loaded progress: {len(self.completed_pages)} pages done")
            except (json.JSONDecodeError, IOError, ValueError, TypeError):
                print("Progress file corrupted, starting fresh")

    def save(self):
        with self._lock:
            data = {
                "metadata": self.metadata,
                "completed_pages": sorted(self.completed_pages),
                "translations": self.translations,
            }
            progress_path = Path(self.progress_file)
            progress_path.parent.mkdir(parents=True, exist_ok=True)
            tmp_path = None
            with tempfile.NamedTemporaryFile(
                "w",
                encoding="utf-8",
                dir=str(progress_path.parent),
                prefix=progress_path.name + ".",
                suffix=".tmp",
                delete=False,
            ) as f:
                tmp_path = Path(f.name)
                json.dump(data, f, ensure_ascii=False, indent=2)
                f.write("\n")
            os.replace(tmp_path, progress_path)

    def is_completed(self, page_num: int) -> bool:
        return page_num in self.completed_pages

    def mark_completed(self, page_num: int, translation: str):
        with self._lock:
            self.completed_pages.add(page_num)
            self.translations[str(page_num)] = translation
        self.save()

    def clear_pages(self, page_nums) -> int:
        cleared = 0
        with self._lock:
            for page_num in page_nums:
                page_cleared = False
                if page_num in self.completed_pages:
                    self.completed_pages.remove(page_num)
                    page_cleared = True
                if str(page_num) in self.translations:
                    self.translations.pop(str(page_num), None)
                    page_cleared = True
                if page_cleared:
                    cleared += 1
        if cleared:
            self.save()
        return cleared

    def get_translation(self, page_num: int) -> str:
        return self.translations.get(str(page_num), "")


# ============================================================
# PDF OVERLAY OUTPUT
# ============================================================

class PDFOverlayWriter:
    """Writes translated text over the original PDF, preserving layout."""

    def __init__(self, source_pdf_path: str, output_pdf_path: str):
        self.source_path = source_pdf_path
        self.output_path = output_pdf_path
        self.doc = pymupdf.open(source_pdf_path)

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        self.close()
        return False

    def overlay_page(self, page_num: int, translated_text: str):
        page = self.doc[page_num]
        page_dict = page.get_text("dict", flags=pymupdf.TEXT_PRESERVE_WHITESPACE)
        blocks = page_dict.get("blocks", [])
        text_blocks = [b for b in blocks if b.get("type") == 0]
        if not text_blocks:
            return

        translated_paragraphs = [p.strip() for p in translated_text.split("\n") if p.strip()]
        clean_paragraphs = []
        for p in translated_paragraphs:
            p = re.sub(r"^#{1,4}\s*", "", p)
            p = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", p)
            if p.strip() == "---" or p.startswith("<!--"):
                continue
            if p.strip():
                clean_paragraphs.append(p)

        para_idx = 0
        for block in text_blocks:
            if para_idx >= len(clean_paragraphs):
                break
            bbox = block["bbox"]
            x0, y0, x1, y1 = bbox
            orig_size = 10
            if block.get("lines"):
                spans = block["lines"][0].get("spans", [])
                if spans:
                    orig_size = spans[0].get("size", 10)
            font_size = min(orig_size * 0.85, 9)
            font_size = max(font_size, 6)

            rect = pymupdf.Rect(x0, y0, x1, y1)
            page.draw_rect(rect, color=None, fill=(1, 1, 1))

            text_to_insert = clean_paragraphs[para_idx]
            try:
                avail_width = x1 - x0 - 4
                if avail_width < 20:
                    para_idx += 1
                    continue
                tw = pymupdf.TextWriter(page.rect)
                font = pymupdf.Font("china-s")
                tw.append((x0 + 2, y0 + font_size + 2), text_to_insert, font=font, fontsize=font_size)
                tw.write_text(page)
            except Exception:
                try:
                    page.insert_text((x0 + 2, y0 + font_size + 2), text_to_insert[:100],
                                     fontsize=font_size, fontname="china-s")
                except Exception:
                    pass
            para_idx += 1

    def save(self):
        ensure_output_parent(self.output_path)
        self.doc.save(self.output_path, garbage=4, deflate=True)
        self.close()

    def close(self):
        if self.doc is not None:
            self.doc.close()
            self.doc = None


# ============================================================
# BATCH CONCURRENT TRANSLATOR
# ============================================================

def translate_batch_concurrent(pages_data, translator, tracker, max_workers=4, progress_callback=None):
    results = {}
    completed_count = 0
    total_count = len(pages_data)
    max_workers = max(1, int(max_workers or 1))

    def report(page_num, translation):
        nonlocal completed_count
        completed_count += 1
        if progress_callback:
            progress_callback(page_num, translation or "", completed_count, total_count)

    def translate_one(page_num, text, prev_ctx):
        translation = translator.translate_chunk(text, page_num, prev_ctx)
        if translation and not is_failed_translation(translation):
            tracker.mark_completed(page_num, translation)
        return page_num, translation

    group_size = max_workers
    prev_context = ""

    for group_start in range(0, len(pages_data), group_size):
        group = pages_data[group_start:group_start + group_size]
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {}
            for page_num, text, page_context in group:
                if tracker.is_completed(page_num):
                    translation = tracker.get_translation(page_num)
                    results[page_num] = translation
                    report(page_num, translation)
                    continue
                if not text.strip():
                    tracker.mark_completed(page_num, "")
                    results[page_num] = ""
                    report(page_num, "")
                    continue
                future = executor.submit(translate_one, page_num, text, prev_context or page_context)
                futures[future] = page_num

            for future in as_completed(futures):
                page_num = futures[future]
                try:
                    page_num, translation = future.result()
                except Exception as exc:
                    translation = f"{TRANSLATION_FAILURE_PREFIX} {exc}]"
                    print(f" p{page_num + 1} failed: {exc}", end="", flush=True)
                results[page_num] = translation or ""
                report(page_num, translation)
                if not is_failed_translation(translation):
                    print(f" p{page_num + 1} done", end="", flush=True)

        if group:
            last_page_num = group[-1][0]
            last_translation = results.get(last_page_num, "")
            if last_translation:
                prev_context = last_translation[-300:]
        print()

    return results
def set_section_columns(section, num=2, space_twips=720):
    """
    设置 Word 分栏。
    space_twips=720 约等于 0.5 英寸，可按需要调小到 360。
    """
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)

    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_section_page_layout(section, columns=1):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.25)

    set_section_columns(section, num=columns, space_twips=520)


def _add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)


def _header_title(title: str) -> str:
    clean = re.sub(r"[_]+", " ", title).strip()
    if " - " in clean:
        clean = clean.split(" - ", 1)[1].strip()
    return clean[:32]


def clear_header_footer_part(part):
    element = part._element
    for child in list(element):
        element.remove(child)


def set_running_header_footer(doc, title: str, header_left: str = "绿色三角洲",
                              header_right: Optional[str] = None):
    right_title = header_right.strip() if header_right else _header_title(title)
    left_title = header_left.strip() if header_left else "绿色三角洲"
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        clear_header_footer_part(section.header)

        table = section.header.add_table(rows=1, cols=2, width=Inches(7.4))
        table.autofit = False
        remove_table_borders(table)
        set_cell_width(table.cell(0, 0), Inches(3.2))
        set_cell_width(table.cell(0, 1), Inches(4.2))

        left_para = table.cell(0, 0).paragraphs[0]
        right_para = table.cell(0, 1).paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for para, text in ((left_para, f"// {left_title} //"), (right_para, f"// {right_title} //")):
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            run = para.add_run(text)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(9)

        clear_header_footer_part(section.footer)
        footer_para = section.footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.paragraph_format.space_before = Pt(0)
        footer_para.paragraph_format.space_after = Pt(0)
        _add_page_number(footer_para)


def set_document_base_layout(doc, columns=1, body_font_size=12.0, line_spacing=1.5,
                             h1_size=None, h2_size=None, h3_size=None):
    set_section_page_layout(doc.sections[0], columns=columns)
    body_font_size = float(body_font_size)
    h1_size = float(h1_size) if h1_size else body_font_size + 16
    h2_size = float(h2_size) if h2_size else body_font_size + 8
    h3_size = float(h3_size) if h3_size else body_font_size + 4

    styles = doc.styles

    # 正文
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(body_font_size)
    normal.paragraph_format.first_line_indent = Pt(body_font_size * 2)
    normal.paragraph_format.line_spacing = line_spacing
    normal.paragraph_format.space_after = Pt(max(3, body_font_size / 2))
    # 一级标题
    h1 = styles["Heading 1"]
    h1.font.name = "黑体"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.font.size = Pt(h1_size)
    h1.font.bold = False
    h1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True

    # 二级标题
    h2 = styles["Heading 2"]
    h2.font.name = "黑体"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h2.font.size = Pt(h2_size)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0xD8, 0x00, 0x00)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    # 三级标题
    h3 = styles["Heading 3"]
    h3.font.name = "黑体"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h3.font.size = Pt(h3_size)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    # 项目符号
    if "List Bullet" in styles:
        bullet = styles["List Bullet"]
        bullet.font.name = "宋体"
        bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        bullet.font.size = Pt(body_font_size)
        bullet.paragraph_format.left_indent = Pt(22)
        bullet.paragraph_format.first_line_indent = Pt(-12)
        bullet.paragraph_format.line_spacing = line_spacing
        bullet.paragraph_format.space_after = Pt(4)


def _translation_blocks(translated_pages):
    blocks = []
    for page_num, translation in translated_pages:
        if not translation.strip():
            continue
        chunks = re.split(r"\n\s*\n", translation)
        for chunk in chunks:
            text = _clean_translated_block(chunk.strip())
            if not text or text == "---" or text.startswith("<!--"):
                continue
            blocks.append({"source_page": page_num, "text": text})
    return blocks


def _clean_translated_block(text: str) -> str:
    lines = []
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        line = _clean_decorative_slash_line(line)
        lines.append(line)

    text = "\n".join(lines)
    text = _dedupe_adjacent_repeated_units(text)
    return text.strip()


def _clean_decorative_slash_line(line: str) -> str:
    if line.count("//") < 2:
        return line

    parts = [p.strip(" /") for p in line.split("//") if p.strip(" /")]
    if not parts:
        return line

    unique_parts = []
    for part in parts:
        normalized = re.sub(r"\s+", "", part).lower()
        if unique_parts and normalized == re.sub(r"\s+", "", unique_parts[-1]).lower():
            continue
        unique_parts.append(part)

    if len(unique_parts) == len(parts):
        return line
    return "// " + " / ".join(unique_parts) + " //"


def _dedupe_adjacent_repeated_units(text: str) -> str:
    patterns = [
        r"([“\"][^”\"\n]{2,120}[”\"])(?:\s*[，,、]?\s*\1)+",
        r"(——[^—\n]{2,60})(?:\s+\1)+",
        r"([^。！？!?\n]{2,120}[。！？!?])(?:\s*\1)+",
    ]
    previous = None
    while previous != text:
        previous = text
        for pattern in patterns:
            text = re.sub(pattern, r"\1", text)
    return text


def _visible_text_length(text: str) -> int:
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"\s+", "", text)
    return len(text)


def _is_markdown_heading(text: str) -> bool:
    return bool(re.match(r"^#{1,6}\s+", text.strip()))


def _is_plain_heading_line(text: str) -> bool:
    clean = re.sub(r"\*\*(.+?)\*\*", r"\1", text.strip())
    clean = re.sub(r"\*(.+?)\*", r"\1", clean)
    if clean.startswith(("#", "-", "\u2022", "//", "——", "“", "\"")):
        return False
    visible = re.sub(r"\s+", "", clean)
    if not (2 <= len(visible) <= 18):
        return False
    if re.search(r"[。！？!?；;：:，,、（）()《》\"“”]", clean):
        return False
    if re.search(r"\d", clean):
        return False
    return True


def _format_page_ranges(page_nums):
    nums = sorted({p + 1 for p in page_nums})
    if not nums:
        return ""
    ranges = []
    start = prev = nums[0]
    for num in nums[1:]:
        if num == prev + 1:
            prev = num
            continue
        ranges.append(f"{start}" if start == prev else f"{start}-{prev}")
        start = prev = num
    ranges.append(f"{start}" if start == prev else f"{start}-{prev}")
    return ", ".join(ranges)


def paginate_translated_blocks(translated_pages, min_chars=1000, max_chars=1500,
                               page_layouts: Optional[dict] = None,
                               split_on_layout=False):
    """Group translated Markdown blocks into reading pages without splitting blocks."""
    pages = []
    current = []
    current_len = 0
    current_layout = None

    def flush():
        nonlocal current, current_len, current_layout
        if current:
            pages.append({
                "layout": current_layout or "columns",
                "blocks": current,
            })
            current = []
            current_len = 0
            current_layout = None

    for block in _translation_blocks(translated_pages):
        block_len = _visible_text_length(block["text"])
        starts_heading = _is_markdown_heading(block["text"])
        block_layout = "columns"
        if page_layouts:
            block_layout = page_layouts.get(block["source_page"], "columns")

        if split_on_layout and current and block_layout != current_layout:
            flush()

        if starts_heading and current and current_len >= min_chars:
            flush()
        elif current and current_len + block_len > max_chars and current_len >= min_chars:
            flush()

        if current_layout is None:
            current_layout = block_layout
        current.append(block)
        current_len += block_len

    flush()
    return pages


def write_markdown_output(translated_pages, md_output: str, title: str, toc: str = "",
                          min_chars=1000, max_chars=1500):
    ensure_output_parent(md_output)
    reading_pages = paginate_translated_blocks(translated_pages, min_chars, max_chars)
    with open(md_output, "w", encoding="utf-8") as f:
        f.write(f"# {title} — 中文翻译\n\n---\n\n")

        if toc:
            f.write(toc)
            f.write("\n---\n\n")

        for page_idx, page in enumerate(reading_pages, 1):
            blocks = page["blocks"]
            source_pages = _format_page_ranges([b["source_page"] for b in blocks])
            f.write(f"<!-- Reading Page {page_idx}; Source PDF Pages: {source_pages} -->\n\n")
            for block in blocks:
                f.write(_format_markdown_block(block["text"]))
                f.write("\n\n")
            f.write("---\n\n")


def _html_inline(text: str) -> str:
    text = html.escape(text.strip())
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
    return text


def _is_markdown_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def _html_table(lines: list[str]) -> str:
    rows = []
    for raw_line in lines:
        cells = [cell.strip() for cell in raw_line.strip().strip("|").split("|")]
        if cells:
            rows.append(cells)
    if not rows:
        return ""

    header = rows[0]
    body = rows[2:] if len(rows) > 1 and _is_markdown_table_separator(lines[1]) else rows[1:]
    parts = ['<table class="aid-table">', "<thead><tr>"]
    for cell in header:
        parts.append(f"<th>{_html_inline(cell)}</th>")
    parts.append("</tr></thead><tbody>")
    for row in body:
        parts.append("<tr>")
        for idx in range(len(header)):
            cell = row[idx] if idx < len(row) else ""
            parts.append(f"<td>{_html_inline(cell)}</td>")
        parts.append("</tr>")
    parts.append("</tbody></table>")
    return "".join(parts)


def _html_handout_card(lines: list[str]) -> str:
    parts = ['<div class="handout-card">']
    for idx, line in enumerate(lines):
        clean = line.strip()
        if not clean:
            continue
        if idx == 0 and len(re.sub(r"\s+", "", clean)) <= 80:
            parts.append(f"<h3>{_html_inline(clean)}</h3>")
        else:
            parts.append(f"<p>{_html_inline(clean)}</p>")
    parts.append("</div>")
    return "".join(parts)


def _html_block(text: str) -> str:
    parts = []
    lines = text.split("\n")
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        clean_line = line.strip()
        idx += 1
        if not clean_line or clean_line == "---" or clean_line.startswith("<!--"):
            continue

        if clean_line == "[[TOC]]":
            continue

        if clean_line.startswith("```toc"):
            toc_lines = []
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                toc_lines.append(lines[idx].rstrip())
                idx += 1
            if idx < len(lines) and lines[idx].strip().startswith("```"):
                idx += 1
            parts.append('<pre class="toc-card">' + html.escape("\n".join(toc_lines)) + "</pre>")
            continue

        if clean_line.startswith("|"):
            table_lines = [clean_line]
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].strip())
                idx += 1
            parts.append(_html_table(table_lines))
            continue

        if clean_line.startswith(">"):
            quote_lines = [clean_line.lstrip(">").strip()]
            while idx < len(lines) and lines[idx].strip().startswith(">"):
                quote_lines.append(lines[idx].strip().lstrip(">").strip())
                idx += 1
            parts.append(_html_handout_card(quote_lines))
            continue

        if clean_line.startswith("### "):
            parts.append(f"<h3>{_html_inline(clean_line[4:])}</h3>")
        elif clean_line.startswith("## "):
            parts.append(f"<h2>{_html_inline(clean_line[3:])}</h2>")
        elif clean_line.startswith("# "):
            parts.append(f"<h1>{_html_inline(clean_line[2:])}</h1>")
        elif _is_plain_heading_line(clean_line):
            parts.append(f"<h2>{_html_inline(clean_line)}</h2>")
        elif clean_line.startswith("- ") or clean_line.startswith("\u2022 "):
            parts.append(f"<ul><li>{_html_inline(clean_line[2:])}</li></ul>")
        else:
            parts.append(f"<p>{_html_inline(clean_line)}</p>")
    return "\n".join(parts)


def write_html_output(translated_pages, html_output: str, title: str, subtitle: str = "中文翻译",
                      min_chars=1200, max_chars=1800, columns=2,
                      header_left="绿色三角洲", header_right=None,
                      page_layouts: Optional[dict] = None):
    """Write translated content as a printable dual-column HTML document."""
    min_chars = int(min_chars)
    max_chars = int(max_chars)
    columns = int(columns)
    if min_chars < 1 or max_chars < min_chars:
        raise ValueError("HTML 阅读页字数范围无效")
    if columns not in (1, 2):
        raise ValueError("HTML 正文分栏只支持 1 或 2 栏")

    ensure_output_parent(html_output)
    right_title = html.escape((header_right or _header_title(title)).strip())
    left_title = html.escape((header_left or "绿色三角洲").strip())
    safe_title = html.escape(title)
    safe_subtitle = html.escape(subtitle or "")
    reading_pages = paginate_translated_blocks(
        translated_pages,
        min_chars,
        max_chars,
        page_layouts=page_layouts,
        split_on_layout=True,
    )

    css = f"""
    :root {{
        color-scheme: light;
        --paper: #f7f2e8;
        --ink: #111111;
        --red: #d80000;
        --muted: #77716a;
        --rule: #b9b0a5;
    }}
    * {{ box-sizing: border-box; }}
    body {{
        margin: 0;
        background: #d8d2cc;
        color: var(--ink);
        font-family: "Noto Serif SC", "Songti SC", "SimSun", serif;
        line-height: 1.72;
    }}
    .sheet {{
        width: 8.5in;
        min-height: 11in;
        margin: 18px auto;
        padding: 0.34in 0.48in 0.52in;
        background:
            radial-gradient(circle at 12% 18%, rgba(160, 132, 93, 0.11), transparent 22%),
            radial-gradient(circle at 85% 70%, rgba(126, 96, 62, 0.08), transparent 24%),
            var(--paper);
        box-shadow: 0 4px 18px rgba(0, 0, 0, 0.22);
        break-after: page;
        page-break-after: always;
    }}
    .running-head {{
        display: flex;
        justify-content: space-between;
        gap: 2rem;
        align-items: baseline;
        margin-bottom: 0.28in;
        padding-bottom: 0.08in;
        border-bottom: 1px solid var(--rule);
        color: var(--muted);
        font: 10pt "Courier New", monospace;
        letter-spacing: 0;
    }}
    .running-head span:last-child {{
        text-align: right;
    }}
    .content {{
        column-count: {columns};
        column-gap: 0.52in;
        font-size: 12pt;
    }}
    h1, h2, h3 {{
        break-after: avoid;
        page-break-after: avoid;
        font-family: "Noto Sans SC", "Microsoft YaHei", sans-serif;
        line-height: 1.12;
        letter-spacing: 0;
    }}
    h1 {{
        column-span: all;
        margin: 0 0 0.28in;
        font-size: 26pt;
        font-weight: 500;
    }}
    h2 {{
        margin: 0.06in 0 0.18in;
        color: var(--red);
        font-size: 20pt;
        font-weight: 700;
    }}
    h3 {{
        margin: 0.16in 0 0.08in;
        font-size: 14pt;
        font-weight: 700;
    }}
    p {{
        margin: 0 0 0.11in;
        text-indent: 2em;
    }}
    h1 + p, h2 + p, h3 + p {{
        text-indent: 0;
    }}
    ul {{
        margin: 0 0 0.12in 1.2em;
        padding: 0;
    }}
    li {{
        margin-bottom: 0.04in;
    }}
    .aid-table {{
        column-span: all;
        width: 100%;
        margin: 0.18in 0 0.24in;
        border-collapse: collapse;
        font-family: "Courier New", "VT323", monospace;
        font-size: 9pt;
        line-height: 1.35;
        background: rgba(232, 226, 204, 0.58);
    }}
    .aid-table th,
    .aid-table td {{
        padding: 0.06in 0.08in;
        border-top: 1px dashed #5f574f;
        border-bottom: 1px dashed #5f574f;
        vertical-align: top;
    }}
    .aid-table th {{
        text-transform: uppercase;
        letter-spacing: 0.04em;
        font-weight: 700;
    }}
    .handout .content {{
        column-count: 1;
        max-width: 7.25in;
        margin: 0 auto;
    }}
    .handout-card {{
        column-span: all;
        margin: 0.16in 0 0.26in;
        padding: 0.16in 0.22in;
        background: rgba(244, 225, 125, 0.62);
        border: 1px solid rgba(145, 126, 55, 0.42);
        box-shadow: 0 0.06in 0.12in rgba(0, 0, 0, 0.14);
        font-family: "Courier New", "VT323", monospace;
        break-inside: avoid;
        page-break-inside: avoid;
    }}
    .handout-card h3 {{
        margin: 0 0 0.08in;
        color: var(--ink);
        font-family: "Courier New", "VT323", monospace;
        font-size: 13pt;
        text-transform: uppercase;
    }}
    .handout-card p {{
        margin: 0 0 0.06in;
        text-indent: 0;
        font-size: 10.5pt;
        line-height: 1.5;
    }}
    .toc .content {{
        column-count: 2;
        column-gap: 0.38in;
        font-size: 10pt;
    }}
    .toc h1 {{
        column-span: all;
        margin-bottom: 0.18in;
        color: var(--ink);
        font-family: "Courier New", "VT323", monospace;
        font-size: 24pt;
    }}
    .toc-card {{
        margin: 0;
        white-space: pre-wrap;
        font-family: "Courier New", "VT323", monospace;
        font-size: 8.7pt;
        line-height: 1.22;
        break-inside: avoid;
        page-break-inside: avoid;
    }}
    .source-pages {{
        column-span: all;
        margin-top: 0.24in;
        color: var(--muted);
        font: 8.5pt "Courier New", monospace;
        text-align: right;
    }}
    .cover .content {{
        column-count: 1;
    }}
    .sheet.single .content {{
        column-count: 1;
        max-width: 7.25in;
        margin: 0 auto;
        font-size: 12pt;
        line-height: 1.74;
    }}
    .sheet.single h2 {{
        margin-top: 0.08in;
        padding-bottom: 0.05in;
        border-bottom: 1px solid var(--rule);
    }}
    .sheet.single .source-pages {{
        column-span: none;
    }}
    .cover-title {{
        margin-top: 1.15in;
        font: 32pt "Noto Sans SC", "Microsoft YaHei", sans-serif;
        letter-spacing: 0;
    }}
    .cover-subtitle {{
        color: #2d73b9;
        font: 14pt "Noto Sans SC", "Microsoft YaHei", sans-serif;
        text-indent: 0;
    }}
    @page {{
        size: Letter;
        margin: 0;
    }}
    @media print {{
        body {{ background: white; }}
        .sheet {{
            margin: 0;
            box-shadow: none;
            width: 8.5in;
            min-height: 11in;
        }}
    }}
    """

    chunks = [
        "<!doctype html>",
        '<html lang="zh-CN">',
        "<head>",
        '<meta charset="utf-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1">',
        f"<title>{safe_title} - 中文翻译</title>",
        f"<style>{css}</style>",
        "</head>",
        "<body>",
        '<section class="sheet cover">',
        f'<header class="running-head"><span>// {left_title} //</span><span>// {right_title} //</span></header>',
        '<main class="content">',
        f'<h1 class="cover-title">{safe_title}</h1>',
        f'<p class="cover-subtitle">{safe_subtitle}</p>' if safe_subtitle else "",
        "</main>",
        "</section>",
    ]

    for page_idx, page in enumerate(reading_pages, 1):
        blocks = page["blocks"]
        layout = page.get("layout", "columns")
        source_pages = _format_page_ranges([b["source_page"] for b in blocks])
        chunks.extend([
            f'<section class="sheet {html.escape(layout)}">',
            f'<header class="running-head"><span>// {left_title} //</span><span>// {right_title} //</span></header>',
            '<main class="content">',
        ])
        for block in blocks:
            chunks.append(_html_block(block["text"]))
        chunks.append(f'<div class="source-pages">Reading Page {page_idx}; Source PDF Pages: {html.escape(source_pages)}</div>')
        chunks.extend(["</main>", "</section>"])

    chunks.extend(["</body>", "</html>", ""])
    with open(html_output, "w", encoding="utf-8") as f:
        f.write("\n".join(chunk for chunk in chunks if chunk != ""))


def _format_markdown_block(text: str) -> str:
    lines = []
    for line in text.split("\n"):
        stripped = line.strip()
        if _is_plain_heading_line(stripped):
            lines.append(f"### {stripped}")
        else:
            lines.append(line)
    return "\n".join(lines)


def _write_word_block(doc, text: str):
    for line in text.split("\n"):
        line = line.strip()
        if not line or line == "---" or line.startswith("<!--"):
            continue

        clean_line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        clean_line = re.sub(r"\*(.+?)\*", r"\1", clean_line)

        if clean_line.startswith("### "):
            p = doc.add_heading(clean_line[4:], level=3)
            p.paragraph_format.first_line_indent = Pt(0)
        elif clean_line.startswith("## "):
            p = doc.add_heading(clean_line[3:], level=2)
            p.paragraph_format.first_line_indent = Pt(0)
        elif clean_line.startswith("# "):
            p = doc.add_heading(clean_line[2:], level=1)
            p.paragraph_format.first_line_indent = Pt(0)
        elif clean_line.startswith(">"):
            card_text = clean_line.lstrip(">").strip()
            if not card_text:
                continue
            p = doc.add_paragraph(card_text)
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.right_indent = Pt(8)
            p.paragraph_format.first_line_indent = Pt(0)
            p_pr = p._p.get_or_add_pPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F4E17D")
            p_pr.append(shading)
        elif _is_plain_heading_line(clean_line):
            p = doc.add_heading(clean_line, level=2)
            p.paragraph_format.first_line_indent = Pt(0)
        elif clean_line.startswith("- ") or clean_line.startswith("\u2022 "):
            p = doc.add_paragraph(clean_line[2:], style="List Bullet")
            p.paragraph_format.first_line_indent = Pt(-8)
        else:
            doc.add_paragraph(clean_line)


def write_word_output(translated_pages, docx_output: str, title: str, subtitle: str = "\u4e2d\u6587\u7ffb\u8bd1",
                      min_chars=1000, max_chars=1500, body_font_size=12.0,
                      line_spacing=1.5, columns=2, header_left="绿色三角洲",
                      header_right=None, hard_page_breaks=False):
    """Write translated Markdown-like page content to a Word document."""
    if not HAS_DOCX:
        raise RuntimeError("Word output requires python-docx")
    min_chars = int(min_chars)
    max_chars = int(max_chars)
    columns = int(columns)
    body_font_size = float(body_font_size)
    line_spacing = float(line_spacing)
    if min_chars < 1 or max_chars < min_chars:
        raise ValueError("Word 阅读页字数范围无效")
    if columns not in (1, 2):
        raise ValueError("Word 正文分栏只支持 1 或 2 栏")
    if not 6 <= body_font_size <= 24:
        raise ValueError("Word 正文字号超出支持范围")
    if not 0.8 <= line_spacing <= 3.0:
        raise ValueError("Word 行距超出支持范围")
    ensure_output_parent(docx_output)

    doc = DocxDocument()
    set_document_base_layout(doc, columns=1, body_font_size=body_font_size, line_spacing=line_spacing)

    title_para = doc.add_heading(title.upper(), level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if subtitle:
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.style = doc.styles["Normal"]
        subtitle_para.paragraph_format.first_line_indent = Pt(0)
        if subtitle_para.runs:
            subtitle_para.runs[0].font.color.rgb = RGBColor(0x2D, 0x73, 0xB9)
            subtitle_para.runs[0].font.bold = True
        doc.add_paragraph()

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_section_page_layout(body_section, columns=columns)
    set_running_header_footer(doc, title, header_left=header_left, header_right=header_right)

    reading_pages = paginate_translated_blocks(translated_pages, min_chars, max_chars)
    for page_idx, page in enumerate(reading_pages):
        blocks = page["blocks"]
        if hard_page_breaks and page_idx > 0:
            doc.add_page_break()
        for block in blocks:
            _write_word_block(doc, block["text"])

    doc.save(docx_output)
# ============================================================
# MAIN ORCHESTRATOR
# ============================================================

def translate_pdf(pdf_path, output_path, api_key, glossary_path=None,
                  model="deepseek-v4-pro", start_page=0, end_page=None,
                  output_format="markdown", max_workers=1):
    print("=" * 60)
    print("  DG TRPG PDF Translator v2.0")
    print("=" * 60)
    print()

    if output_format not in SUPPORTED_OUTPUT_FORMATS:
        raise ValueError(f"不支持的输出格式：{output_format}")
    if not pdf_path or not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF 文件不存在：{pdf_path}")
    if glossary_path and not os.path.exists(glossary_path):
        raise FileNotFoundError(f"术语表文件不存在：{glossary_path}")
    max_workers = max(1, min(16, int(max_workers or 1)))
    ensure_output_parent(output_path)

    stats = TokenStats()

    print(f"Opening PDF: {pdf_path}")
    extractor = PDFExtractor(pdf_path)
    try:
        total = extractor.total_pages
        print(f"   Total pages: {total}")

        start_page, end_page = normalize_page_range(start_page, end_page, total)
        print(f"   Range: page {start_page + 1} to {end_page}")
        print(f"   Workers: {max_workers}")
        print(f"   Format: {output_format}")
        print()

        glossary = {}
        if glossary_path:
            print(f"Loading glossary: {glossary_path}")
            glossary = load_glossary(glossary_path)
            print(f"   Loaded {len(glossary)} terms")
            print()

        print(f"Engine: DeepSeek V4 ({model})")
        translator = Translator(api_key=api_key, model=model, stats=stats)
        translator.set_glossary(glossary)
        print()

        progress_file = output_path + ".progress.json"
        progress_metadata = build_progress_metadata(
            pdf_path=pdf_path,
            glossary_path=glossary_path,
            model=model,
            start_page=start_page,
            end_page=end_page,
        )
        tracker = ProgressTracker(progress_file, expected_metadata=progress_metadata)
        if tracker.metadata_mismatches:
            print("⚠️  进度文件与当前设置不一致。")
            for mismatch in tracker.metadata_mismatches[:5]:
                print(f"   - {mismatch}")
            if tracker.ignored_existing_progress:
                print("   已保留文件但本次不复用旧译文。")
        print()

        print("Extracting text and analyzing chapters...")
        pages_text = {}
        page_layouts = {}
        for page_num in range(start_page, end_page):
            page_layouts[page_num] = extractor.detect_page_layout(page_num)
            text = extractor.extract_page(page_num)
            pages_text[page_num] = text

        extractor.finalize_chapters()
        toc = extractor.chapter_detector.get_toc_markdown()
        if toc:
            print(f"   Detected {len(extractor.chapter_detector.headings)} headings")
        else:
            print("   No clear chapter structure detected")
        print()

        print("Translating...")
        print("-" * 40)
        start_time = time.time()

        if max_workers > 1:
            print(f"   Concurrent mode: {max_workers} workers")
            pages_data = []
            prev_text = ""
            for page_num in range(start_page, end_page):
                text = pages_text.get(page_num, "")
                context = prev_text[-300:] if prev_text else ""
                pages_data.append((page_num, text, context))
                context_text = extractor.get_context_text(page_num)
                if context_text.strip():
                    prev_text = context_text
            results = translate_batch_concurrent(pages_data, translator, tracker, max_workers)
            translated_pages = [(pn, t) for pn, t in results.items() if t.strip()]
        else:
            translated_pages = []
            prev_translation_tail = ""
            pages_to_process = list(range(start_page, end_page))
            total_to_do = len(pages_to_process)
            done_count = 0

            for page_num in pages_to_process:
                done_count += 1
                progress_pct = done_count / total_to_do * 100

                if tracker.is_completed(page_num):
                    translation = tracker.get_translation(page_num)
                    if translation:
                        translated_pages.append((page_num, translation))
                        prev_translation_tail = translation[-300:]
                    print(f"  [skip] Page {page_num + 1}/{total}")
                    continue

                text = pages_text.get(page_num, "")
                if not text.strip():
                    print(f"  [empty] Page {page_num + 1}/{total}")
                    tracker.mark_completed(page_num, "")
                    continue

                print(f"  [{progress_pct:.0f}%] Page {page_num + 1}/{total}", end="", flush=True)
                translation = translator.translate_chunk(text, page_num, prev_context=prev_translation_tail)

                if translation and not is_failed_translation(translation):
                    translated_pages.append((page_num, translation))
                    tracker.mark_completed(page_num, translation)
                    prev_translation_tail = translation[-300:]
                    print(f" done (Y{stats.cost_yuan:.3f})")
                elif is_failed_translation(translation):
                    translated_pages.append((page_num, translation))
                    print(" failed; not cached")
                else:
                    print(f" empty result")
                    tracker.mark_completed(page_num, "")
                time.sleep(0.3)

        elapsed = time.time() - start_time
        print("-" * 40)
        print()

        translated_pages_sorted = sorted(translated_pages, key=lambda x: x[0])
        failed_pages = [
            page_num + 1
            for page_num, translation in translated_pages_sorted
            if is_failed_translation(translation)
        ]
        if failed_pages:
            print("⚠️  以下页翻译失败且未写入进度缓存: " + ", ".join(map(str, failed_pages[:20])))

        # Determine output base name (without extension)
        output_base = output_path
        for ext in (".md", ".pdf", ".docx"):
            if output_base.endswith(ext):
                output_base = output_base[:-len(ext)]
                break

        if glossary:
            report_output = output_base + "_glossary_report.md"
            print(f"  生成术语命中报告: {report_output}")
            write_glossary_report(pages_text, glossary, report_output, Path(pdf_path).stem)
            print("   ✓ 术语报告输出完成")

        # HTML output
        if output_format in ("html", "both", "all"):
            html_output = output_base + ".html"
            print(f"  生成 HTML: {html_output}")
            try:
                write_html_output(
                    translated_pages_sorted,
                    html_output,
                    Path(pdf_path).stem,
                    page_layouts=page_layouts,
                )
                print("   ✅ HTML 输出完成")
            except Exception as e:
                print(f"   ❌ HTML 输出失败: {e}")

        # Markdown output
        if output_format in ("markdown", "both", "all"):
            md_output = output_base + ".md"
            print(f"  生成 Markdown: {md_output}")
            write_markdown_output(translated_pages_sorted, md_output, Path(pdf_path).stem, toc)
            print("   ✅ Markdown 输出完成")

        # Word output
        if output_format in ("word", "all"):
            if not HAS_DOCX:
                print("  ⚠️  Word 输出需要 python-docx，请运行: pip install python-docx")
            else:
                docx_output = output_base + ".docx"
                print(f"  生成 Word 文档: {docx_output}")
                try:
                    write_word_output(translated_pages_sorted, docx_output, Path(pdf_path).stem)
                    print("   ✓ Word 输出完成")

                except Exception as e:
                    print(f"   ❌ Word 输出失败: {e}")

        page_count = len([t for _, t in translated_pages_sorted if t.strip()])
        print(f"\n  共翻译 {page_count} 页")

        print()
        print(f"Time: {elapsed:.1f}s ({elapsed/60:.1f} min)")
        print()
        print(stats.summary())
        print()
        print(f"Progress file: {progress_file}")
        print(f"Output: {output_path}")

    finally:
        extractor.close()


# ============================================================
# CLI with Config File Support
# ============================================================

def load_config(config_path: str) -> dict:
    """Load configuration from a JSON file."""
    if not os.path.exists(config_path):
        print(f"❌ 配置文件不存在: {config_path}")
        sys.exit(1)
    try:
        with open(config_path, "r", encoding="utf-8") as f:
            config = json.load(f)
    except json.JSONDecodeError as exc:
        print(f"❌ 配置文件 JSON 格式错误: {exc}")
        sys.exit(1)
    if not isinstance(config, dict):
        print("❌ 配置文件顶层必须是 JSON 对象。")
        sys.exit(1)
    return config


def main():
    parser = argparse.ArgumentParser(
        description="绿色三角洲 PDF 翻译工具 v2.0",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  # 使用配置文件（推荐，最简单）
  python translate_pdf.py --config config.json

  # 命令行参数
  python translate_pdf.py "THE MILLENNIUM.pdf" --api-key sk-xxx --format html --workers 4

  # 指定术语表和范围
  python translate_pdf.py "THE MILLENNIUM.pdf" --api-key sk-xxx \
      --glossary glossary.tsv --start 0 --end 5

  # 使用更便宜的模型
  python translate_pdf.py "THE MILLENNIUM.pdf" --api-key sk-xxx \
      --model deepseek-v4-flash --workers 8
        """
    )

    parser.add_argument("pdf", nargs="?", default=None, help="输入 PDF 文件路径")
    parser.add_argument("--config", "-c", default=None, help="配置文件路径（JSON）")
    parser.add_argument("--api-key", default=None, help="DeepSeek API Key")
    parser.add_argument("--output", "-o", default=None, help="输出文件路径（不含扩展名）")
    parser.add_argument("--glossary", "-g", default=None, help="术语表文件路径（TSV 格式）")
    parser.add_argument("--model", default=None, help="模型名称（默认: deepseek-v4-pro）")
    parser.add_argument("--format", "-f", choices=["markdown", "html", "word", "both", "all"],
                        default=None, help="输出格式: markdown/html/word/both/all（默认: markdown）")
    parser.add_argument("--workers", "-w", type=int, default=None,
                        help="并发线程数（默认: 1，推荐: 4）")
    parser.add_argument("--start", type=int, default=None, help="起始页码（从0开始）")
    parser.add_argument("--end", type=int, default=None, help="结束页码（不含）")

    args = parser.parse_args()

    # Load config file if specified
    config = {}
    if args.config:
        config = load_config(args.config)

    # Merge: command line args override config file
    pdf_path = args.pdf or config.get("pdf")
    api_key = args.api_key or config.get("api_key")
    output_path = args.output or config.get("output")
    glossary_path = args.glossary or config.get("glossary")
    model = args.model or config.get("model", "deepseek-v4-pro")
    output_format = args.format or config.get("format", "markdown")
    workers = args.workers if args.workers is not None else config.get("workers", 1)
    start_page = args.start if args.start is not None else config.get("start", 0)
    end_page = args.end if args.end is not None else config.get("end")

    # Validate required params
    if not pdf_path:
        print("❌ 缺少 PDF 文件路径。请通过参数或配置文件指定。")
        parser.print_help()
        sys.exit(1)
    if not api_key:
        print("❌ 缺少 API Key。请通过 --api-key 或配置文件指定。")
        sys.exit(1)
    if output_format not in SUPPORTED_OUTPUT_FORMATS:
        print(f"❌ 不支持的输出格式: {output_format}")
        sys.exit(1)

    # Default output path
    if output_path is None:
        pdf_stem = Path(pdf_path).stem
        output_path = f"{pdf_stem}_cn.md"

    # Validate
    if not os.path.exists(pdf_path):
        print(f"❌ PDF 文件不存在: {pdf_path}")
        sys.exit(1)
    if glossary_path and not os.path.exists(glossary_path):
        print(f"❌ 术语表文件不存在: {glossary_path}")
        sys.exit(1)

    try:
        workers = int(workers)
        start_page = int(start_page or 0)
        end_page = None if end_page is None or end_page == "" else int(end_page)
    except (TypeError, ValueError):
        print("❌ workers/start/end 必须是整数。")
        sys.exit(1)

    if workers < 1:
        workers = 1
    elif workers > 16:
        print("⚠️  并发数上限为 16，已自动调整")
        workers = 16

    # Run
    translate_pdf(
        pdf_path=pdf_path,
        output_path=output_path,
        api_key=api_key,
        glossary_path=glossary_path,
        model=model,
        start_page=start_page,
        end_page=end_page,
        output_format=output_format,
        max_workers=workers,
    )


if __name__ == "__main__":
    main()

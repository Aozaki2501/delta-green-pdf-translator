#!/usr/bin/env python3
"""
Delta Green PDF Translator — Web UI (Streamlit)
"""
import streamlit as st
import os
import re
import time
from pathlib import Path
from translate_pdf import (
    PDFExtractor, Translator, ProgressTracker, TokenStats,
    PDFOverlayWriter, load_glossary, translate_batch_concurrent,
    set_document_base_layout
)
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# === UI THEME ===
st.set_page_config(
    page_title="DG Translator Terminal",
    page_icon="🖧",
    layout="wide",
)

st.markdown("""
<style>
    /* 引入复古等宽字体 */
    @import url('https://fonts.googleapis.com/css2?family=Courier+Prime:ital,wght@0,400;0,700;1,400&family=VT323&display=swap');

    :root {
        --term-bg: #050505;
        --term-panel: #0a0c0a;
        --term-green: #33ff33;
        --term-dark-green: #114411;
        --term-text: #c8d6c8;
        --term-border: #1f3b22;
        --term-alert: #ff3333;
    }

    /* 强制覆盖全局背景和字体 */
    .stApp, .stAppHeader {
        background-color: var(--term-bg) !important;
        color: var(--term-text) !important;
        font-family: "Courier Prime", "Courier New", monospace !important;
    }

    /* 侧边栏样式 */
    section[data-testid="stSidebar"] {
        background-color: var(--term-panel) !important;
        border-right: 2px solid var(--term-border) !important;
    }
    
    /* 标题样式：更具复古电子感 */
    h1, h2, h3, .hero-title {
        color: var(--term-green) !important;
        font-family: "VT323", "Courier New", monospace !important;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        font-weight: normal;
    }

    p, label, .stMarkdown {
        color: var(--term-text) !important;
        font-size: 0.95rem;
    }

    /* 顶层 Hero 面板 */
    .hero {
        background-color: transparent;
        border: 1px solid var(--term-border);
        border-left: 6px solid var(--term-green);
        padding: 20px 24px;
        margin-bottom: 30px;
    }

    .hero-title {
        font-size: 2.2rem;
        margin-bottom: 8px;
    }

    .hero-subtitle {
        color: var(--term-text);
        opacity: 0.8;
    }

    /* 输入框与选择器样式 - 去除圆角，方正硬朗 */
    .stTextInput input,
    .stNumberInput input,
    .stSelectbox div[data-baseweb="select"],
    .stMultiSelect div[data-baseweb="select"] {
        background-color: var(--term-panel) !important;
        border: 1px solid var(--term-border) !important;
        border-radius: 0px !important;
        color: var(--term-green) !important;
        font-family: "Courier Prime", monospace !important;
    }

    /* 上传组件虚线框 */
    div[data-testid="stFileUploader"] {
        background-color: var(--term-panel) !important;
        border: 1px dashed var(--term-border) !important;
        border-radius: 0px !important;
        padding: 16px;
    }

    /* 主按钮样式：终端高亮效果 */
    .stButton>button {
        background-color: var(--term-bg) !important;
        color: var(--term-green) !important;
        border: 1px solid var(--term-green) !important;
        border-radius: 0px !important;
        height: 48px;
        font-weight: bold;
        letter-spacing: 0.1em;
        transition: all 0.2s ease;
        text-transform: uppercase;
    }

    .stButton>button:hover {
        background-color: var(--term-green) !important;
        color: var(--term-bg) !important;
    }

    /* 进度条变绿 */
    .stProgress > div > div > div {
        background-color: var(--term-green) !important;
    }

    /* 状态信息卡片 */
    div[data-testid="stMetric"] {
        background-color: var(--term-panel) !important;
        border: 1px solid var(--term-border) !important;
        border-radius: 0px !important;
        padding: 15px;
    }
    
    div[data-testid="stMetricValue"] {
        color: var(--term-green) !important;
    }

    /* 消除基础结构过大的留白 */
    .block-container {
        max-width: 1200px;
        padding-top: 2rem;
    }
</style>
""", unsafe_allow_html=True)

# === HEADER ===
st.markdown("""
<div class="hero">
    <div class="hero-title">RESTRICTED // DG-TRANSLATOR-SYS</div>
    <div class="hero-subtitle">
        > INITIALIZING AUTOMATED TRANSLATION PROTOCOL...<br>
        > PURPOSE: EXTRACT, TRANSLATE AND COMPILE ENCRYPTED PDF DATA.<br>
        > STATUS: READY FOR OPERATOR INPUT.
    </div>
</div>
""", unsafe_allow_html=True)
st.markdown("""
<div class="hero">
    <div class="hero-title">Delta Green PDF Translator</div>
    <div class="hero-subtitle">
        将英文 TRPG PDF 提取、翻译并输出为 Markdown、Word 或实验性排版 PDF。
        更适合先产出可校对译稿，再做人工排版。
    </div>
</div>
""", unsafe_allow_html=True)
with st.sidebar:
    st.header("翻译配置")

    api_key = st.text_input("API Key", type="password", placeholder="sk-...")
    model = st.selectbox("模型", ["deepseek-v4-pro", "deepseek-v4-flash"])
    workers = st.slider("并发线程", 1, 16, 4)

    formats = st.multiselect(
        "输出格式",
        ["Markdown", "PDF", "Word"],
        default=["Word"],
    )

    start_page = st.number_input("起始页", value=0, min_value=0)
    end_page_str = st.text_input("结束页", value="", placeholder="留空表示全部")

# === MAIN ===
st.markdown('<div class="section-card">', unsafe_allow_html=True)

st.subheader("输入文件")
st.caption("上传原始 PDF。术语表为可选 TSV / TXT / CSV 文件。")

col1, col2 = st.columns([1.2, 1])
with col1:
    pdf_file = st.file_uploader("PDF 文件", type=["pdf"], label_visibility="collapsed")
with col2:
    glossary_file = st.file_uploader("术语表，可选", type=["tsv", "txt", "csv"], label_visibility="collapsed")

st.markdown("</div>", unsafe_allow_html=True)

if st.button("🔺 开始翻译", type="primary", use_container_width=True):
    if not pdf_file:
        st.error("✗ 请上传 PDF 文件")
    elif not api_key:
        st.error("✗ 请输入 API Key")
    else:
        # Create organized directories
        upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
        os.makedirs(upload_dir, exist_ok=True)
        os.makedirs(output_dir, exist_ok=True)

        # Save uploaded files to uploads/
        pdf_path = os.path.join(upload_dir, pdf_file.name)
        with open(pdf_path, "wb") as f:
            f.write(pdf_file.read())

        glossary_path = None
        if glossary_file:
            glossary_path = os.path.join(upload_dir, glossary_file.name)
            with open(glossary_path, "wb") as f:
                f.write(glossary_file.read())

        end_page = int(end_page_str) if end_page_str.strip() else None
        output_base = os.path.join(output_dir, f"{Path(pdf_file.name).stem}_cn")

        # Init
        stats = TokenStats()
        extractor = PDFExtractor(pdf_path)
        total = extractor.total_pages
        if end_page is None or end_page > total:
            end_page = total

        glossary = load_glossary(glossary_path) if glossary_path else {}
        translator = Translator(api_key=api_key, model=model, stats=stats)
        translator.set_glossary(glossary)

        progress_file = output_base + ".progress.json"
        tracker = ProgressTracker(progress_file)

        # Extract
        st.info(f"📑 提取文本: {total} 页, 翻译 {start_page+1}-{end_page} 页")
        pages_text = {}
        for pn in range(start_page, end_page):
            pages_text[pn] = extractor.extract_page(pn)
        extractor.finalize_chapters()
        toc = extractor.chapter_detector.get_toc_markdown()

        # Translate
        progress_bar = st.progress(0)
        status_text = st.empty()
        translated_pages = []
        prev_tail = ""
        pages_list = list(range(start_page, end_page))
        total_to_do = len(pages_list)

        for idx, pn in enumerate(pages_list):
            pct = (idx + 1) / total_to_do
            progress_bar.progress(pct)
            status_text.text(f"翻译中: 第 {pn+1} 页 [{pct*100:.0f}%] | ¥{stats.cost_yuan:.3f}")

            if tracker.is_completed(pn):
                t = tracker.get_translation(pn)
                if t:
                    translated_pages.append((pn, t))
                    prev_tail = t[-300:]
                continue

            text = pages_text.get(pn, "")
            if not text.strip():
                tracker.mark_completed(pn, "")
                continue

            translation = translator.translate_chunk(text, pn, prev_context=prev_tail)
            if translation:
                translated_pages.append((pn, translation))
                tracker.mark_completed(pn, translation)
                prev_tail = translation[-300:]
            time.sleep(0.2)

        progress_bar.progress(1.0)
        status_text.text("✓ 翻译完成!")
        translated_pages_sorted = sorted(translated_pages, key=lambda x: x[0])

        # Stats
        col_a, col_b, col_c = st.columns(3)
        col_a.metric("📄 页数", f"{len(translated_pages_sorted)}")
        col_b.metric("💰 费用", f"¥{stats.cost_yuan:.3f}")
        col_c.metric("🔢 Token", f"{stats.total_tokens:,}")

                # Output & Download
        if "Markdown" in formats:
            md_path = output_base + ".md"
            pdf_title = Path(pdf_file.name).stem
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(f"# {pdf_title} — 中文翻译\n\n---\n\n")

                if toc:
                    f.write(toc)
                    f.write("\n---\n\n")

                for pn, t in translated_pages_sorted:
                    if t.strip():
                        f.write(f"<!-- Page {pn + 1} -->\n\n")
                        f.write(t)
                        f.write("\n\n---\n\n")

            with open(md_path, "rb") as f:
                st.download_button(
                    "📥 下载 Markdown",
                    f,
                    file_name=f"{Path(pdf_file.name).stem}_cn.md",
                )

        if "PDF" in formats:
            pdf_out_path = output_base + ".pdf"
            try:
                writer = PDFOverlayWriter(pdf_path, pdf_out_path)

                for pn, t in translated_pages_sorted:
                    if t.strip():
                        writer.overlay_page(pn, t)

                writer.save()

                with open(pdf_out_path, "rb") as f:
                    st.download_button(
                        "📥 下载 PDF",
                        f,
                        file_name=f"{Path(pdf_file.name).stem}_cn.pdf",
                    )
            except Exception as e:
                st.error(f"PDF 输出失败：{e}")

        if "Word" in formats:
            if not HAS_DOCX:
                st.warning("Word 输出需要 python-docx，请运行：pip install python-docx")
            else:
                docx_path = output_base + ".docx"
                doc = DocxDocument()
                set_document_base_layout(doc)

                pdf_title = Path(pdf_file.name).stem
                title_para = doc.add_heading(pdf_title.upper(), level=1)
                title_para.alignment = 0

                for pn, t in translated_pages_sorted:
                    if not t.strip():
                        continue

                    for line in t.split("\n"):
                        line = line.strip()

                        if not line or line == "---" or line.startswith("<!--"):
                            continue

                        if line.startswith("### "):
                            doc.add_heading(line[4:], level=3)
                        elif line.startswith("## "):
                            doc.add_heading(line[3:], level=2)
                        elif line.startswith("# "):
                            doc.add_heading(line[2:], level=1)
                        elif line.startswith("- ") or line.startswith("• "):
                            doc.add_paragraph(line[2:], style="List Bullet")
                        else:
                            clean_line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                            clean_line = re.sub(r"\*(.+?)\*", r"\1", clean_line)
                            doc.add_paragraph(clean_line)

                doc.save(docx_path)

                with open(docx_path, "rb") as f:
                    st.download_button(
                        "📥 下载 Word",
                        f,
                        file_name=f"{Path(pdf_file.name).stem}_cn.docx",
                    )

        extractor.close()
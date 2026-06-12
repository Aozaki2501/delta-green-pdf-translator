"""
Word/DOCX exporter module.

Provides write_word_output and all Word-specific helper functions for
generating translated DOCX documents with proper layout, columns,
headers/footers, and card formatting.

Dependencies: python-docx (optional), exporters._shared
"""

import re
import os
from typing import Optional

from exporters._shared import (
    _is_plain_heading_line,
    _is_soft_subheading_line,
    _is_markdown_table_separator_row,
    _looks_like_stat_block,
    paginate_translated_blocks,
    _layout_uses_columns,
    _normalize_heading_markup,
    _normalize_marker_line,
    _display_title,
    _header_title as _shared_header_title,
    _without_image_blocks,
    _looks_like_markdown_table_row,
    _collect_strict_markdown_table,
    _strip_single_cell_pipe_fragment,
    _strip_list_marker,
    _strip_quote_prefix,
)
from core.utils import ensure_output_parent

# ---------------------------------------------------------------------------
# Optional python-docx import
# ---------------------------------------------------------------------------

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ---------------------------------------------------------------------------
# Word layout helpers
# ---------------------------------------------------------------------------

def set_section_columns(section, num=2, space_twips=720):
    """
    设置 Word 分栏。
    space_twips=720 约等于 0.5 英寸，可按需要调小到 360。
    """
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)

    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_table_borders(table, color="B8B8B8", val="dashed", size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), val)
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_section_page_layout(section, columns=1):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.25)

    set_section_columns(section, num=columns, space_twips=520)


def _columns_for_layout(layout: str, default_columns: int) -> int:
    if layout == "toc":
        return 2
    return int(default_columns) if _layout_uses_columns(layout) else 1


def _image_asset_path(asset) -> str:
    if isinstance(asset, dict):
        return str(asset.get("path") or "")
    return str(asset or "")


def _image_asset_placement(asset) -> str:
    if isinstance(asset, dict):
        placement = str(asset.get("placement") or "full").lower()
        if placement in {"left", "right", "full"}:
            return placement
    return "full"


def _add_word_field(paragraph, instruction: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)


def _add_page_number(paragraph):
    _add_word_field(paragraph, "PAGE")


def _header_title(title: str) -> str:
    return _shared_header_title(title)


def clear_header_footer_part(part):
    element = part._element
    for child in list(element):
        element.remove(child)


def _clear_section_chrome(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_header_footer_part(section.header)
    clear_header_footer_part(section.footer)


def _restart_page_numbering(section, start: int = 1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(int(start)))


def _continue_page_numbering(section):
    pg_num = section._sectPr.find(qn("w:pgNumType"))
    if pg_num is not None:
        section._sectPr.remove(pg_num)


def _set_section_header(section, left_title: str, right_title: str):
    clear_header_footer_part(section.header)
    table = section.header.add_table(rows=1, cols=2, width=Inches(7.4))
    table.autofit = False
    remove_table_borders(table)
    set_cell_width(table.cell(0, 0), Inches(3.2))
    set_cell_width(table.cell(0, 1), Inches(4.2))

    left_para = table.cell(0, 0).paragraphs[0]
    right_para = table.cell(0, 1).paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for para in (left_para, right_para):
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0

    run = left_para.add_run(f"// {left_title} //")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)

    run = right_para.add_run(f"// {right_title} //")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)


def _set_section_footer(section, source_note: str = "", show_page_number: bool = True):
    clear_header_footer_part(section.footer)
    table = section.footer.add_table(rows=1, cols=2, width=Inches(7.4))
    table.autofit = False
    remove_table_borders(table)
    set_cell_width(table.cell(0, 0), Inches(4.8))
    set_cell_width(table.cell(0, 1), Inches(2.6))

    left_para = table.cell(0, 0).paragraphs[0]
    right_para = table.cell(0, 1).paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for para in (left_para, right_para):
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0

    if source_note:
        run = left_para.add_run(source_note)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(9)

    if show_page_number:
        _add_page_number(right_para)


def set_section_header_footer(section, title: str, header_left: str = "绿色三角洲",
                              header_right: Optional[str] = None, source_note: str = "",
                              show_header: bool = True, show_page_number: bool = True,
                              restart_page_number: Optional[int] = None):
    _clear_section_chrome(section)
    if restart_page_number is not None:
        _restart_page_numbering(section, restart_page_number)
    if show_header:
        right_title = header_right.strip() if header_right else _header_title(title)
        left_title = header_left.strip() if header_left else "绿色三角洲"
        _set_section_header(section, left_title, right_title)
    _set_section_footer(section, source_note=source_note, show_page_number=show_page_number)


def set_document_base_layout(doc, columns=1, body_font_size=12.0, line_spacing=1.5,
                             h1_size=None, h2_size=None, h3_size=None, h4_size=None):
    set_section_page_layout(doc.sections[0], columns=columns)
    body_font_size = float(body_font_size)
    h1_size = float(h1_size) if h1_size else 20
    h2_size = float(h2_size) if h2_size else 18
    h3_size = float(h3_size) if h3_size else 16
    h4_size = float(h4_size) if h4_size else 16

    styles = doc.styles

    # 正文
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(body_font_size)
    normal.paragraph_format.first_line_indent = Pt(body_font_size * 2)
    normal.paragraph_format.line_spacing = line_spacing
    normal.paragraph_format.space_after = Pt(max(3, body_font_size / 2))
    # 一级标题
    h1 = styles["Heading 1"]
    h1.font.name = "黑体"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.font.size = Pt(h1_size)
    h1.font.bold = False
    h1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True

    # 二级标题
    h2 = styles["Heading 2"]
    h2.font.name = "黑体"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h2.font.size = Pt(h2_size)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0xD8, 0x00, 0x00)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    # 三级标题
    h3 = styles["Heading 3"]
    h3.font.name = "黑体"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h3.font.size = Pt(h3_size)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    # 四级标题
    h4 = styles["Heading 4"]
    h4.font.name = "黑体"
    h4._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h4.font.size = Pt(h4_size)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(4)
    h4.paragraph_format.keep_with_next = True

    # 项目符号
    if "List Bullet" in styles:
        bullet = styles["List Bullet"]
        bullet.font.name = "宋体"
        bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        bullet.font.size = Pt(body_font_size)
        bullet.paragraph_format.left_indent = Pt(22)
        bullet.paragraph_format.first_line_indent = Pt(-12)
        bullet.paragraph_format.line_spacing = line_spacing
        bullet.paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# Word content writing helpers
# ---------------------------------------------------------------------------

def _is_table_line(line: str) -> bool:
    """Check if a line is part of a Markdown table."""
    return _looks_like_markdown_table_row(line)


def _table_cells(line: str) -> list[str]:
    stripped = _strip_quote_prefix(line).strip().strip("|")
    cells = [cell.strip() for cell in stripped.split("|")]
    return [re.sub(r"^#{1,6}\s*", "", cell).strip() for cell in cells]


def _plain_text(line: str) -> str:
    clean = _strip_quote_prefix(line)
    clean = _normalize_heading_markup(clean)
    clean = _normalize_marker_line(clean)
    clean = _strip_single_cell_pipe_fragment(clean)
    clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)
    clean = re.sub(r"\*(.+?)\*", r"\1", clean)
    return clean.strip()


def _split_toc_entry(line: str) -> tuple[str, str] | None:
    match = re.match(r"^(?P<title>.*?)\s*(?:[.\-]{3,}|\s{2,})\s*(?P<page>\d{1,4})\s*$", line.strip())
    if not match:
        return None
    title = re.sub(r"[.\-]{3,}\s*$", "", match.group("title")).strip(" -\t")
    page = match.group("page").strip()
    if not title:
        return None
    return title, page


def _split_card_segments(text: str):
    """Split block text into segments: normal (dual-column), card (single-column),
    and table (single-column). This allows the Word renderer to switch column
    layout around cards and tables."""
    segments = []
    normal_lines = []
    card_lines = []
    stat_lines = []
    image_lines = []
    full_title_lines = []
    table_lines = []
    quote_lines = []
    toc_lines = []
    in_card = False
    in_stat = False
    in_image = False
    in_full_title = False
    in_quote = False
    in_toc = False

    def flush_normal():
        nonlocal normal_lines
        if any(line.strip() for line in normal_lines):
            segments.append(("normal", "\n".join(normal_lines).strip()))
        normal_lines = []

    def flush_card():
        nonlocal card_lines
        if any(line.strip() for line in card_lines):
            segments.append(("card", "\n".join(card_lines).strip()))
        card_lines = []

    def flush_stat():
        nonlocal stat_lines
        if any(line.strip() for line in stat_lines):
            segments.append(("stat", "\n".join(stat_lines).strip()))
        stat_lines = []

    def flush_image():
        nonlocal image_lines
        if any(line.strip() for line in image_lines):
            segments.append(("image", "\n".join(image_lines).strip()))
        image_lines = []

    def flush_full_title():
        nonlocal full_title_lines
        if any(line.strip() for line in full_title_lines):
            segments.append(("full_title", "\n".join(full_title_lines).strip()))
        full_title_lines = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            segments.append(("table", "\n".join(table_lines).strip()))
        table_lines = []

    def flush_quote():
        nonlocal quote_lines
        if any(line.strip() for line in quote_lines):
            segments.append(("card", "\n".join(quote_lines).strip()))
        quote_lines = []

    def flush_toc():
        nonlocal toc_lines
        if any(line.strip() for line in toc_lines):
            segments.append(("toc", "\n".join(toc_lines).strip()))
        toc_lines = []

    raw_lines = text.splitlines()
    line_index = 0
    while line_index < len(raw_lines):
        raw_line = raw_lines[line_index]
        line = _normalize_marker_line(raw_line)
        stripped = line.strip()
        if in_quote and not stripped:
            quote_lines.append("")
            line_index += 1
            continue
        if in_quote and stripped.startswith(">"):
            quote_lines.append(_strip_quote_prefix(raw_line))
            line_index += 1
            continue
        if in_quote:
            flush_quote()
            in_quote = False

        if line == "[FULL_WIDTH_TITLE]":
            flush_normal()
            flush_table()
            flush_card()
            flush_stat()
            flush_image()
            flush_quote()
            flush_toc()
            in_full_title = True
            line_index += 1
            continue
        if line == "[/FULL_WIDTH_TITLE]":
            flush_full_title()
            in_full_title = False
            line_index += 1
            continue
        if line == "[CARD]":
            flush_normal()
            flush_table()
            flush_quote()
            flush_toc()
            in_card = True
            line_index += 1
            continue
        if line == "[/CARD]":
            flush_card()
            in_card = False
            line_index += 1
            continue
        if line == "[STAT_BLOCK]":
            flush_normal()
            flush_table()
            flush_quote()
            flush_toc()
            in_stat = True
            line_index += 1
            continue
        if line == "[/STAT_BLOCK]":
            flush_stat()
            in_stat = False
            line_index += 1
            continue
        if line == "[IMAGE]":
            flush_normal()
            flush_table()
            flush_quote()
            flush_toc()
            in_image = True
            line_index += 1
            continue
        if line == "[/IMAGE]":
            flush_image()
            in_image = False
            line_index += 1
            continue
        if stripped.startswith("```toc"):
            flush_normal()
            flush_table()
            flush_quote()
            in_toc = True
            line_index += 1
            continue
        if in_toc and stripped.startswith("```"):
            flush_toc()
            in_toc = False
            line_index += 1
            continue
        if in_card:
            card_lines.append(raw_line)
            line_index += 1
            continue
        if in_stat:
            stat_lines.append(raw_line)
            line_index += 1
            continue
        if in_image:
            image_lines.append(raw_line)
            line_index += 1
            continue
        if in_full_title:
            full_title_lines.append(raw_line)
            line_index += 1
            continue
        if in_toc:
            toc_lines.append(raw_line)
            line_index += 1
            continue

        if stripped.startswith(">"):
            flush_normal()
            flush_table()
            flush_toc()
            in_quote = True
            quote_lines.append(_strip_quote_prefix(raw_line))
            line_index += 1
            continue

        table_candidate, next_index = _collect_strict_markdown_table(
            raw_lines,
            line_index,
            lambda value: _normalize_marker_line(str(value)),
        )
        if table_candidate:
            flush_normal()
            table_lines.extend(table_candidate)
            flush_table()
            line_index = next_index
            continue

        if table_lines:
            flush_table()
        normal_lines.append(raw_line)
        line_index += 1

    if in_full_title:
        flush_full_title()
    elif in_card:
        flush_card()
    elif in_stat:
        flush_stat()
    elif in_image:
        flush_image()
    elif in_quote:
        flush_quote()
    elif in_toc:
        flush_toc()
    else:
        flush_table()
        flush_normal()
    return segments


def _write_word_toc(doc, text: str):
    for raw_line in text.split("\n"):
        line = _plain_text(raw_line)
        if not line:
            continue
        entry = _split_toc_entry(line)
        if entry:
            title, page = entry
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                Inches(3.45),
                WD_TAB_ALIGNMENT.RIGHT,
                WD_TAB_LEADER.DOTS,
            )
            title_run = paragraph.add_run(title)
            title_run.font.name = "Courier New"
            title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            title_run.font.size = Pt(8.5)
            page_run = paragraph.add_run(f"\t{page}")
            page_run.font.name = "Courier New"
            page_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            page_run.font.size = Pt(8.5)
            continue

        paragraph = doc.add_paragraph(line)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        if paragraph.runs:
            run = paragraph.runs[0]
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.size = Pt(10)
            run.font.bold = True


def _write_word_block(doc, text: str, layout: str = "columns"):
    plain_indent = layout == "columns"
    centered = layout in {"credits", "art"}

    def tune_paragraph(paragraph):
        if not plain_indent:
            paragraph.paragraph_format.first_line_indent = Pt(0)
        if centered:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line or line == "---" or line.startswith("<!--"):
            continue
        clean_line = _plain_text(raw_line)
        if not clean_line:
            continue

        if clean_line == "[[TOC]]":
            continue

        if layout == "toc":
            _write_word_toc(doc, re.sub(r"^#{1,6}\s*", "", clean_line))
            continue

        if clean_line.startswith("#### "):
            p = doc.add_heading(clean_line[5:], level=4)
            p.paragraph_format.first_line_indent = Pt(0)
            tune_paragraph(p)
        elif clean_line.startswith("### "):
            p = doc.add_heading(clean_line[4:], level=3)
            p.paragraph_format.first_line_indent = Pt(0)
            tune_paragraph(p)
        elif clean_line.startswith("## "):
            p = doc.add_heading(clean_line[3:], level=2)
            p.paragraph_format.first_line_indent = Pt(0)
            tune_paragraph(p)
        elif clean_line.startswith("# "):
            p = doc.add_heading(clean_line[2:], level=1)
            p.paragraph_format.first_line_indent = Pt(0)
            tune_paragraph(p)
        elif _is_plain_heading_line(clean_line):
            p = doc.add_heading(clean_line, level=2)
            p.paragraph_format.first_line_indent = Pt(0)
            tune_paragraph(p)
        elif _is_soft_subheading_line(clean_line):
            p = doc.add_heading(clean_line, level=4)
            p.paragraph_format.first_line_indent = Pt(0)
            tune_paragraph(p)
        elif (list_item := _strip_list_marker(clean_line)):
            p = doc.add_paragraph(list_item, style="List Bullet")
            p.paragraph_format.first_line_indent = Pt(-8)
            tune_paragraph(p)
        else:
            p = doc.add_paragraph(clean_line)
            tune_paragraph(p)


def _write_word_table(doc, text: str):
    """Render a Markdown table as a Word table."""
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if len(lines) < 2 or not _is_markdown_table_separator_row(lines[1]):
        # Not a strict Markdown table, fall back to plain text.
        for line in lines:
            doc.add_paragraph(_strip_quote_prefix(line).strip("| "))
        return

    # Parse header
    header_cells = _table_cells(lines[0])

    # Skip separator line (e.g. |---|---|---|)
    data_start = 1
    if len(lines) > 1 and _is_markdown_table_separator_row(lines[1]):
        data_start = 2

    # Parse data rows
    data_rows = []
    for line in lines[data_start:]:
        cells = _table_cells(line)
        data_rows.append(cells)

    col_count = len(header_cells)
    if col_count < 1:
        return

    # Create Word table
    table = doc.add_table(rows=1 + len(data_rows), cols=col_count)
    table.style = "Table Grid"

    # Write header
    for idx, cell_text in enumerate(header_cells):
        if idx < col_count:
            cell = table.rows[0].cells[idx]
            cell.text = cell_text
            # Bold header
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.bold = True

    # Write data rows
    for row_idx, row_cells in enumerate(data_rows):
        for col_idx in range(col_count):
            cell_text = row_cells[col_idx] if col_idx < len(row_cells) else ""
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)

    # Add spacing after table
    doc.add_paragraph()


def _set_paragraph_runs(paragraph, *, font_name="宋体", font_size=10.5, bold=False):
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(font_size)
        if bold:
            run.bold = True


def _style_card_paragraph(paragraph, *, title=False, subheading=False, list_item=False):
    paragraph.paragraph_format.left_indent = Pt(20 if list_item else 12)
    paragraph.paragraph_format.right_indent = Pt(12)
    paragraph.paragraph_format.first_line_indent = Pt(-12 if list_item else 0)
    paragraph.paragraph_format.space_before = Pt(6 if title else 0)
    paragraph.paragraph_format.space_after = Pt(3)
    _set_paragraph_left_rule(paragraph)
    _set_paragraph_shading(paragraph)
    _set_paragraph_runs(
        paragraph,
        font_size=11 if title else (10.8 if subheading else 10.3),
        bold=title or subheading,
    )


def _set_paragraph_left_rule(paragraph, color="B0891C"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "14")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)


def _set_paragraph_shading(paragraph, fill="FFF7D6"):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _write_word_card(doc, text: str):
    clean_lines = [_plain_text(line) for line in text.split("\n")]
    clean_lines = [line for line in clean_lines if line]
    if not clean_lines:
        return

    idx = 0
    first = re.sub(r"^#{1,6}\s*", "", clean_lines[0]).strip()
    if len(re.sub(r"\s+", "", first)) <= 80:
        p = doc.add_paragraph(first)
        _style_card_paragraph(p, title=True)
        idx = 1

    while idx < len(clean_lines):
        clean_line = clean_lines[idx]
        table_lines, next_idx = _collect_strict_markdown_table(clean_lines, idx)
        if table_lines:
            _write_word_table(doc, "\n".join(table_lines))
            idx = next_idx
            continue

        list_item = _strip_list_marker(clean_line)
        if list_item:
            items = [list_item]
            idx += 1
            while idx < len(clean_lines):
                peek = clean_lines[idx]
                next_item = _strip_list_marker(peek)
                if next_item:
                    items.append(next_item)
                    idx += 1
                    continue
                break
            for item in items:
                p = doc.add_paragraph(item, style="List Bullet")
                _style_card_paragraph(p, list_item=True)
            continue

        if re.match(r"^#{1,6}\s+", clean_line) or _is_soft_subheading_line(clean_line):
            p = doc.add_paragraph(re.sub(r"^#{1,6}\s*", "", clean_line).strip())
            _style_card_paragraph(p, subheading=True)
            idx += 1
            continue

        p = doc.add_paragraph(re.sub(r"^#{1,6}\s*", "", clean_line).strip())
        _style_card_paragraph(p)
        idx += 1


def _write_word_stat_block(doc, text: str):
    if not _looks_like_stat_block(text):
        _write_word_card(doc, text)
        return
    for idx, line in enumerate(text.split("\n")):
        clean_line = line.strip()
        if not clean_line:
            continue
        clean_line = re.sub(r"\*\*(.+?)\*\*", r"\1", clean_line)
        clean_line = re.sub(r"\*(.+?)\*", r"\1", clean_line)
        p = doc.add_paragraph(clean_line)
        p.paragraph_format.left_indent = Pt(10)
        p.paragraph_format.right_indent = Pt(10)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(5 if idx == 0 else 0)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            if idx == 0:
                run.bold = True
        p_pr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "8")
        top.set(qn("w:space"), "2")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "2")
        borders.append(top)
        borders.append(bottom)
        p_pr.append(borders)


def _write_word_image_placeholder(doc, text: str, image_path="", placement: str = "full"):
    label = " ".join(line.strip() for line in text.splitlines() if line.strip()) or "插图"
    if label.lower() == "illustration placeholder":
        label = "插图"
    asset_path = _image_asset_path(image_path)
    placement = placement if placement in {"left", "right", "full"} else "full"
    if asset_path:
        if not os.path.exists(asset_path):
            raise FileNotFoundError(f"图片资源不存在：{image_path}")
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        if placement == "left":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            image_width = Inches(3.05)
        elif placement == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            image_width = Inches(3.05)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_width = Inches(6.6)
        p.add_run().add_picture(asset_path, width=image_width)
        caption = doc.add_paragraph(label)
        if placement == "left":
            caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif placement == "right":
            caption.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.first_line_indent = Pt(0)
        caption.paragraph_format.space_before = Pt(2)
        caption.paragraph_format.space_after = Pt(8)
        _set_paragraph_runs(caption, font_name="宋体", font_size=9)
        return
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    row = table.rows[0]
    row.height = Inches(1.35)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    set_cell_width(cell, Inches(7.15))
    para = cell.paragraphs[0]
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.add_run(" ")
    caption = doc.add_paragraph(label)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)
    _set_paragraph_runs(caption, font_name="宋体", font_size=9)


def _write_word_full_title(doc, text: str):
    lines = [
        re.sub(r"^\s*#{1,6}\s*", "", line).strip()
        for line in text.splitlines()
        if line.strip()
    ]
    lines = [re.sub(r"\s+", " ", line) for line in lines if line]
    if not lines:
        return
    p = doc.add_heading(lines[0], level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6 if len(lines) > 1 else 18)
    if len(lines) > 1:
        subtitle = doc.add_paragraph(" ".join(lines[1:]))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.first_line_indent = Pt(0)
        subtitle.paragraph_format.space_after = Pt(18)


def _with_missing_image_placeholders(translated_pages, source_pages_text=None):
    if not source_pages_text:
        return translated_pages
    patched = []
    placeholder = "[IMAGE]\nIllustration placeholder\n[/IMAGE]"
    for page_num, text in translated_pages:
        source_text = source_pages_text.get(page_num, "")
        missing_count = source_text.count("[IMAGE]") - str(text).count("[IMAGE]")
        if missing_count > 0:
            extra = "\n\n".join(placeholder for _ in range(missing_count))
            text = f"{text.rstrip()}\n\n{extra}"
        patched.append((page_num, text))
    return patched


# ---------------------------------------------------------------------------
# Main Word output function
# ---------------------------------------------------------------------------

def write_word_output(translated_pages, docx_output: str, title: str, subtitle: str = "中文翻译",
                      min_chars=1000, max_chars=1500, body_font_size=12.0,
                      line_spacing=1.5, columns=2, header_left="绿色三角洲",
                      header_right=None, hard_page_breaks=False,
                      source_pages_text=None, source_page_labels: Optional[dict] = None,
                      page_layouts=None,
                      image_assets: Optional[dict] = None):
    """Write translated Markdown-like page content to a Word document."""
    if not HAS_DOCX:
        raise RuntimeError("Word output requires python-docx")
    min_chars = int(min_chars)
    max_chars = int(max_chars)
    columns = int(columns)
    body_font_size = float(body_font_size)
    line_spacing = float(line_spacing)
    if min_chars < 1 or max_chars < min_chars:
        raise ValueError("Word 阅读页字数范围无效")
    if columns not in (1, 2):
        raise ValueError("Word 正文分栏只支持 1 或 2 栏")
    if not 6 <= body_font_size <= 24:
        raise ValueError("Word 正文字号超出支持范围")
    if not 0.8 <= line_spacing <= 3.0:
        raise ValueError("Word 行距超出支持范围")
    ensure_output_parent(docx_output)
    translated_pages = _without_image_blocks(translated_pages)
    reading_pages = paginate_translated_blocks(
        translated_pages,
        min_chars,
        max_chars,
        page_layouts=page_layouts,
        split_on_layout=True,
    )
    display_title = _display_title(title, reading_pages)

    doc = DocxDocument()
    set_document_base_layout(doc, columns=1, body_font_size=body_font_size, line_spacing=line_spacing)
    set_section_header_footer(doc.sections[0], display_title, show_header=False, show_page_number=False)

    title_para = doc.add_heading(display_title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.first_line_indent = Pt(0)
    title_para.paragraph_format.space_before = Pt(120)
    title_para.paragraph_format.space_after = Pt(18)

    if subtitle:
        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.style = doc.styles["Normal"]
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_para.paragraph_format.first_line_indent = Pt(0)
        subtitle_para.paragraph_format.space_after = Pt(0)
        if subtitle_para.runs:
            subtitle_para.runs[0].font.color.rgb = RGBColor(0x2D, 0x73, 0xB9)
            subtitle_para.runs[0].font.bold = True
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)

    if not reading_pages:
        doc.save(docx_output)
        return

    first_layout = reading_pages[0].get("layout", "columns")
    current_page_columns = _columns_for_layout(first_layout, columns)
    current_show_header = first_layout != "toc"
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_page_layout(body_section, columns=current_page_columns)
    set_section_header_footer(
        body_section,
        display_title,
        header_left=header_left,
        header_right=header_right,
        show_header=current_show_header,
        show_page_number=True,
        restart_page_number=1,
    )

    for page_idx, page in enumerate(reading_pages):
        blocks = page["blocks"]
        layout = page.get("layout", "columns")
        page_columns = _columns_for_layout(layout, columns)
        show_header = layout != "toc"
        if hard_page_breaks and page_idx > 0:
            doc.add_page_break()
        if page_idx > 0 and (page_columns != current_page_columns or show_header != current_show_header):
            body_section = doc.add_section(WD_SECTION.CONTINUOUS)
            _continue_page_numbering(body_section)
            set_section_page_layout(body_section, columns=page_columns)
            set_section_header_footer(
                body_section,
                display_title,
                header_left=header_left,
                header_right=header_right,
                show_header=show_header,
                show_page_number=True,
            )
        current_page_columns = page_columns
        current_show_header = show_header
        for block in blocks:
            for kind, content in _split_card_segments(block["text"]):
                if kind == "card":
                    _write_word_card(doc, content)
                elif kind == "stat":
                    _write_word_stat_block(doc, content)
                elif kind == "image":
                    continue
                elif kind == "toc":
                    _write_word_toc(doc, content)
                elif kind == "full_title":
                    title_section = doc.add_section(WD_SECTION.CONTINUOUS)
                    _continue_page_numbering(title_section)
                    set_section_page_layout(title_section, columns=1)
                    title_section.header.is_linked_to_previous = True
                    title_section.footer.is_linked_to_previous = True
                    _write_word_full_title(doc, content)
                    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
                    _continue_page_numbering(body_section)
                    set_section_page_layout(body_section, columns=page_columns)
                    body_section.header.is_linked_to_previous = True
                    body_section.footer.is_linked_to_previous = True
                elif kind == "table":
                    table_section = doc.add_section(WD_SECTION.CONTINUOUS)
                    _continue_page_numbering(table_section)
                    set_section_page_layout(table_section, columns=1)
                    table_section.header.is_linked_to_previous = True
                    table_section.footer.is_linked_to_previous = True
                    _write_word_table(doc, content)
                    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
                    _continue_page_numbering(body_section)
                    set_section_page_layout(body_section, columns=page_columns)
                    body_section.header.is_linked_to_previous = True
                    body_section.footer.is_linked_to_previous = True
                else:
                    _write_word_block(doc, content, layout=layout)

    doc.save(docx_output)

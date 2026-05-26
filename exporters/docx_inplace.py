"""
Word/DOCX in-place replacement exporter.

Copies the original .docx file and replaces text content with translations
while preserving all formatting (styles, fonts, colors, tables, text boxes).

Core strategy: copy original file → open copy → replace text runs → save.

Dependencies: python-docx, lxml, shutil
"""

import re
import os
import shutil
import tempfile
from pathlib import Path
from typing import Optional

from core.docx_extractor import DocxBlock, RunMeta, parse_marked_translation

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# Chinese font fallback list (in preference order)
CJK_FONT_FALLBACKS = [
    "Microsoft YaHei",
    "微软雅黑",
    "SimSun",
    "宋体",
    "PingFang SC",
    "Noto Sans CJK SC",
    "Source Han Sans SC",
]

# Fonts known to NOT support CJK characters
NON_CJK_FONTS = {
    "Times New Roman", "Arial", "Calibri", "Cambria", "Georgia",
    "Verdana", "Helvetica", "Courier New", "Consolas", "Palatino",
    "Garamond", "Book Antiqua", "Century Schoolbook",
}


def write_docx_inplace(blocks: list[DocxBlock], translations: dict[int, str],
                        source_path: str, output_path: str,
                        cjk_font: str = "Microsoft YaHei") -> str:
    """
    Write translated Word document preserving original formatting.

    Args:
        blocks: Full list of DocxBlock from DocxExtractor
        translations: dict mapping block index -> translated text
        source_path: Path to original .docx file
        output_path: Where to write the translated .docx
        cjk_font: Chinese font to use when original font doesn't support CJK

    Returns:
        The output file path.
    """
    if not HAS_DOCX:
        raise ImportError("python-docx is required. Run: pip install python-docx")

    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = None

    try:
        with tempfile.NamedTemporaryFile(
            suffix=".docx",
            prefix=out_path.stem + ".",
            dir=str(out_path.parent),
            delete=False,
        ) as tmp_file:
            tmp_path = Path(tmp_file.name)

        shutil.copy2(source_path, tmp_path)
        doc = DocxDocument(str(tmp_path))

        missing_replacements = []
        for block in blocks:
            translated = translations.get(block.index)
            if not translated or not block.translatable:
                continue

            replaced = False
            if block.block_type == "paragraph":
                replaced = _replace_body_paragraph(doc, block, translated, cjk_font)
            elif block.block_type == "table_cell":
                replaced = _replace_table_cell(doc, block, translated, cjk_font)
            elif block.block_type == "textbox_para":
                replaced = _replace_textbox_paragraph(doc, block, translated, cjk_font)
            elif block.block_type in ("header", "footer"):
                replaced = _replace_header_footer(doc, block, translated, cjk_font)

            if not replaced:
                missing_replacements.append(block.index)

        if missing_replacements:
            sample = ", ".join(map(str, missing_replacements[:10]))
            raise RuntimeError(f"Word 写回失败，定位不到 {len(missing_replacements)} 个译文块：{sample}")

        doc.save(str(tmp_path))
        os.replace(tmp_path, out_path)
        tmp_path = None
    finally:
        if tmp_path and tmp_path.exists():
            tmp_path.unlink()
    return str(out_path)


# ============================================================
# PARAGRAPH REPLACEMENT
# ============================================================

def _replace_body_paragraph(doc, block: DocxBlock, translated: str, cjk_font: str):
    """Replace text in a body paragraph."""
    # Parse parent_path like "body.para[5]"
    m = re.match(r'body\.para\[(\d+)\]', block.parent_path)
    if not m:
        return False

    para_idx = int(m.group(1))
    if para_idx >= len(doc.paragraphs):
        return False

    para = doc.paragraphs[para_idx]
    _replace_paragraph_text(para, block, translated, cjk_font)
    return True


def _replace_table_cell(doc, block: DocxBlock, translated: str, cjk_font: str):
    """Replace text in a table cell paragraph."""
    # Parse parent_path like "table[2].row[1].cell[0].para[0]"
    m = re.match(
        r'table\[(\d+)\]\.row\[(\d+)\]\.cell\[(\d+)\]\.para\[(\d+)\]',
        block.parent_path
    )
    if not m:
        return False

    table_idx = int(m.group(1))
    row_idx = int(m.group(2))
    cell_idx = int(m.group(3))
    para_idx = int(m.group(4))

    if table_idx >= len(doc.tables):
        return False
    table = doc.tables[table_idx]
    if row_idx >= len(table.rows):
        return False
    row = table.rows[row_idx]
    if cell_idx >= len(row.cells):
        return False
    cell = row.cells[cell_idx]
    if para_idx >= len(cell.paragraphs):
        return False
    para = cell.paragraphs[para_idx]

    _replace_paragraph_text(para, block, translated, cjk_font)
    return True


def _replace_textbox_paragraph(doc, block: DocxBlock, translated: str, cjk_font: str):
    """Replace text in a text box paragraph via XML manipulation."""
    # Find all txbxContent elements in document body
    body = doc.element.body
    txbx_elements = body.findall('.//' + qn('w:txbxContent'))

    if block.textbox_index >= len(txbx_elements):
        return False

    txbx = txbx_elements[block.textbox_index]
    para_elements = txbx.findall(qn('w:p'))

    if block.para_in_textbox >= len(para_elements):
        return False

    p_elem = para_elements[block.para_in_textbox]
    _replace_p_element_text(p_elem, block, translated, cjk_font)
    return True


def _replace_header_footer(doc, block: DocxBlock, translated: str, cjk_font: str):
    """Replace text in header/footer paragraphs."""
    # Parse parent_path like "section[0].header.para[0]"
    m = re.match(
        r'section\[(\d+)\]\.(header|footer|first_page_header|first_page_footer|even_page_header|even_page_footer)\.para(?:\[(\d+)\])?',
        block.parent_path,
    )
    if not m:
        return False

    section_idx = int(m.group(1))
    hf_type = m.group(2)
    para_idx = int(m.group(3)) if m.group(3) is not None else None

    if section_idx >= len(doc.sections):
        return False

    section = doc.sections[section_idx]
    containers = {
        "header": section.header,
        "footer": section.footer,
        "first_page_header": section.first_page_header,
        "first_page_footer": section.first_page_footer,
        "even_page_header": section.even_page_header,
        "even_page_footer": section.even_page_footer,
    }
    container = containers[hf_type]
    if not container:
        return False

    if para_idx is not None and para_idx < len(container.paragraphs):
        _replace_paragraph_text(container.paragraphs[para_idx], block, translated, cjk_font)
        return True

    # Fallback for old progress paths without paragraph index.
    for para in container.paragraphs:
        if para.text.strip() == block.text:
            _replace_paragraph_text(para, block, translated, cjk_font)
            return True
    return False


# ============================================================
# CORE TEXT REPLACEMENT LOGIC
# ============================================================

def _replace_paragraph_text(para, block: DocxBlock, translated: str, cjk_font: str):
    """
    Replace paragraph text while preserving formatting.

    Strategy:
    - Single run: replace text, keep all formatting
    - Multiple runs with mixed formatting: use marker-based approach
    - Fallback: clear all runs, write as single run with first run's format
    """
    if not para.runs:
        # No runs — add one
        run = para.add_run(translated)
        _apply_cjk_font(run, cjk_font)
        return

    # Check if translation has format markers
    has_markers = bool(re.search(r'<(b|i|bi)>', translated))

    if has_markers and block.runs:
        # Parse markers and create new runs
        new_runs = parse_marked_translation(translated, block.runs)
        _clear_and_write_runs(para, new_runs, cjk_font)
    elif len(para.runs) == 1:
        # Single run: just replace text
        para.runs[0].text = translated
        _apply_cjk_font_to_run(para.runs[0], cjk_font)
    else:
        # Multiple runs, no markers: use first run's format for all
        _clear_and_write_single(para, translated, cjk_font)


def _replace_p_element_text(p_elem, block: DocxBlock, translated: str, cjk_font: str):
    """Replace text in a raw w:p XML element (for text boxes)."""
    # Remove existing runs
    for r_elem in p_elem.findall(qn('w:r')):
        p_elem.remove(r_elem)

    # Check if translation has format markers
    has_markers = bool(re.search(r'<(b|i|bi)>', translated))

    if has_markers and block.runs:
        new_runs = parse_marked_translation(translated, block.runs)
        for run_meta in new_runs:
            r_elem = _create_run_element(run_meta, cjk_font)
            p_elem.append(r_elem)
    else:
        # Single run
        run_meta = RunMeta(text=translated)
        if block.runs:
            run_meta.font_name = block.runs[0].font_name
            run_meta.font_size = block.runs[0].font_size
            run_meta.bold = block.runs[0].bold
            run_meta.italic = block.runs[0].italic
            run_meta.color = block.runs[0].color
        r_elem = _create_run_element(run_meta, cjk_font)
        p_elem.append(r_elem)


def _clear_and_write_runs(para, new_runs: list[RunMeta], cjk_font: str):
    """Clear paragraph runs and write new ones with formatting."""
    # Preserve paragraph properties
    pPr = para._p.find(qn('w:pPr'))

    # Remove all existing runs
    for run in para.runs:
        para._p.remove(run._r)

    # Add new runs
    for run_meta in new_runs:
        if not run_meta.text:
            continue
        run = para.add_run(run_meta.text)
        if run_meta.bold:
            run.bold = True
        if run_meta.italic:
            run.italic = True
        if run_meta.underline:
            run.underline = True
        if run_meta.font_name:
            run.font.name = run_meta.font_name
        if run_meta.font_size:
            run.font.size = Pt(run_meta.font_size)
        if run_meta.color:
            try:
                run.font.color.rgb = RGBColor.from_string(run_meta.color)
            except (ValueError, AttributeError):
                pass
        _apply_cjk_font_to_run(run, cjk_font)


def _clear_and_write_single(para, text: str, cjk_font: str):
    """Clear all runs and write as single run preserving first run's format."""
    if not para.runs:
        para.add_run(text)
        return

    # Capture first run's formatting
    first_run = para.runs[0]
    bold = first_run.bold
    italic = first_run.italic
    underline = first_run.underline
    font_name = first_run.font.name if first_run.font else None
    font_size = first_run.font.size if first_run.font else None
    try:
        color = first_run.font.color.rgb if first_run.font and first_run.font.color else None
    except Exception:
        color = None

    # Clear existing runs
    for run in list(para.runs):
        para._p.remove(run._r)

    # Write new single run
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if underline:
        run.underline = True
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size
    if color:
        try:
            run.font.color.rgb = color
        except Exception:
            pass
    _apply_cjk_font_to_run(run, cjk_font)


# ============================================================
# CJK FONT HELPERS
# ============================================================

def _apply_cjk_font(run, cjk_font: str):
    """Apply CJK font to a python-docx run object."""
    _apply_cjk_font_to_run(run, cjk_font)


def _apply_cjk_font_to_run(run, cjk_font: str):
    """
    Set East Asian font for a run so CJK characters render correctly.
    Only modifies the East Asian font slot; preserves Latin font.
    """
    # Check if current font needs CJK fallback
    current_font = run.font.name if run.font else None
    if current_font and current_font in NON_CJK_FONTS:
        # Set East Asian font family
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), cjk_font)
    elif not current_font:
        # No font set at all — set both
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), cjk_font)


def _create_run_element(run_meta: RunMeta, cjk_font: str):
    """Create a w:r XML element with formatting (for text box replacement)."""
    r_elem = OxmlElement('w:r')

    # Run properties
    rPr = OxmlElement('w:rPr')
    r_elem.append(rPr)

    if run_meta.bold:
        b_elem = OxmlElement('w:b')
        rPr.append(b_elem)

    if run_meta.italic:
        i_elem = OxmlElement('w:i')
        rPr.append(i_elem)

    if run_meta.font_name or cjk_font:
        rFonts = OxmlElement('w:rFonts')
        if run_meta.font_name:
            rFonts.set(qn('w:ascii'), run_meta.font_name)
            rFonts.set(qn('w:hAnsi'), run_meta.font_name)
        rFonts.set(qn('w:eastAsia'), cjk_font)
        rPr.append(rFonts)

    if run_meta.font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(run_meta.font_size * 2)))  # points to half-points
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(run_meta.font_size * 2)))
        rPr.append(szCs)

    if run_meta.color:
        color_elem = OxmlElement('w:color')
        color_elem.set(qn('w:val'), run_meta.color)
        rPr.append(color_elem)

    # Text element
    t_elem = OxmlElement('w:t')
    t_elem.set(qn('xml:space'), 'preserve')
    t_elem.text = run_meta.text
    r_elem.append(t_elem)

    return r_elem

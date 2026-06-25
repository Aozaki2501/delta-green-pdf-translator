"""Word review package helpers."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


@dataclass
class WordReviewItem:
    page_num: int | None
    category: str
    title: str
    detail: str = ""


def build_word_review_items(
    *,
    quality_report=None,
    glossary_candidates=None,
    rule_symbol_issues=None,
    timeline_events=None,
) -> list[WordReviewItem]:
    items: list[WordReviewItem] = []
    if quality_report:
        for issue in getattr(quality_report, "issues", []) or []:
            items.append(WordReviewItem(
                page_num=int(getattr(issue, "page_num", 0) or 0) or None,
                category="问题页索引",
                title=str(getattr(issue, "message", "") or "质量问题"),
                detail=str(getattr(issue, "detail", "") or ""),
            ))
    for candidate in list(glossary_candidates or [])[:30]:
        pages = ", ".join(str(page + 1) for page in getattr(candidate, "pages", [])[:8])
        items.append(WordReviewItem(
            page_num=None,
            category="术语疑点",
            title=str(getattr(candidate, "term", "") or ""),
            detail=f"出现 {getattr(candidate, 'count', 0)} 次；页码：{pages}",
        ))
    for issue in list(rule_symbol_issues or [])[:80]:
        items.append(WordReviewItem(
            page_num=int(getattr(issue, "page_num", 0) or 0) or None,
            category="规则符号",
            title=f"{getattr(issue, 'kind', '')}：{getattr(issue, 'symbol', '')}",
            detail=str(getattr(issue, "message", "") or ""),
        ))
    for event in list(timeline_events or [])[:80]:
        items.append(WordReviewItem(
            page_num=int(getattr(event, "page_num", 0) or 0) or None,
            category="时间线",
            title=str(getattr(event, "marker", "") or ""),
            detail=str(getattr(event, "event", "") or ""),
        ))
    items.sort(key=lambda item: (item.page_num is None, item.page_num or 0, item.category, item.title))
    return items


def render_word_review_markdown(items: list[WordReviewItem], title: str = "") -> str:
    heading = f"# {title} — Word 校对包" if title else "# Word 校对包"
    lines = [
        heading,
        "",
        f"- 校对项：{len(items)}",
        f"- 涉及页数：{len({item.page_num for item in items if item.page_num})}",
        "",
        "## 校对索引",
        "",
    ]
    if not items:
        lines.append("- 暂无。")
    else:
        lines.extend(["| 源页 | 类别 | 标题 | 说明 |", "| ---: | --- | --- | --- |"])
        for item in items:
            page = str(item.page_num) if item.page_num else "-"
            lines.append(
                f"| {page} | {_escape_table(item.category)} | "
                f"{_escape_table(item.title)} | {_escape_table(item.detail)} |"
            )
    lines.append("")
    return "\n".join(lines)


def write_word_review_markdown(items: list[WordReviewItem], output_path: str, title: str = "") -> None:
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(render_word_review_markdown(items, title))


def write_word_review_docx(items: list[WordReviewItem], output_path: str, title: str = "") -> None:
    if not HAS_DOCX:
        raise RuntimeError("Word review docx requires python-docx")
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc = DocxDocument()
    doc.add_heading(f"{title or '文档'} — Word 校对包", level=1)
    intro = doc.add_paragraph(f"校对项：{len(items)}")
    intro.paragraph_format.space_after = Pt(8)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ("源页", "类别", "标题", "说明")
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for item in items:
        cells = table.add_row().cells
        cells[0].text = str(item.page_num) if item.page_num else "-"
        cells[1].text = item.category
        cells[2].text = item.title
        cells[3].text = item.detail
    doc.save(output_path)


def _escape_table(text: str) -> str:
    return str(text or "").replace("|", "\\|").replace("\n", " ")


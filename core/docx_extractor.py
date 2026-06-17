"""
Word/DOCX text extraction with style metadata preservation.

Parses a .docx file into ordered translation units (DocxBlock) that
capture both the text content and the formatting metadata needed to
write back a translated copy with identical visual layout.

Handles: paragraphs, tables (cell-by-cell), text boxes (via XML),
headers/footers (optional).

Dependencies: python-docx, lxml
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


_TABLE_ROLL_MARKER_RE = re.compile(r"^\s*(?:\d{1,3}|[dD]\d+|[ivxlcdmIVXLCDM]+)(?:\s*[-–—]\s*(?:\d{1,3}|[ivxlcdmIVXLCDM]+))?[\.)、:]?\s*$")


def _is_table_roll_marker(text: str) -> bool:
    """Return True for narrow dice/table labels such as 1, 2, 1-2, or d10."""
    return bool(_TABLE_ROLL_MARKER_RE.fullmatch(text or ""))


# ============================================================
# DATA MODEL
# ============================================================

@dataclass
class RunMeta:
    """Format metadata for a single Word run."""
    text: str
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    font_name: Optional[str] = None
    font_size: Optional[float] = None  # in points
    color: Optional[str] = None  # hex color string like "FF0000"
    strike: Optional[bool] = None


@dataclass
class DocxBlock:
    """A single translation unit extracted from a Word document."""
    index: int
    block_type: str  # paragraph, table_cell, textbox_para, header, footer
    text: str  # Full plain text of the block
    translatable: bool
    style_name: str = ""
    runs: list[RunMeta] = field(default_factory=list)
    parent_path: str = ""  # Locator path like "body.para[5]" or "table[2].row[1].cell[0].para[0]"
    # For table cells
    table_index: int = -1
    row_index: int = -1
    cell_index: int = -1
    para_in_cell: int = 0
    # For text boxes
    textbox_index: int = -1
    para_in_textbox: int = 0


# ============================================================
# DOCX EXTRACTOR
# ============================================================

class DocxExtractor:
    """Extracts translation blocks from a Word document with format metadata."""

    def __init__(self, docx_path: str, translate_headers: bool = False):
        if not HAS_DOCX:
            raise ImportError("python-docx is required. Run: pip install python-docx")
        self.docx_path = Path(docx_path)
        self.translate_headers = translate_headers
        self.blocks: list[DocxBlock] = []
        self._doc = None

    def extract(self) -> list[DocxBlock]:
        """Parse the document and return ordered blocks."""
        self._doc = DocxDocument(str(self.docx_path))
        self.blocks = []
        block_index = 0

        # Track document body elements in order (paragraphs and tables interleaved)
        body = self._doc.element.body
        para_count = 0
        table_count = 0

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                # Regular paragraph
                para = self._doc.paragraphs[para_count] if para_count < len(self._doc.paragraphs) else None
                if para is not None:
                    block = self._extract_paragraph(para, block_index, f"body.para[{para_count}]")
                    if block:
                        self.blocks.append(block)
                        block_index += 1

                    # Check for text boxes inside this paragraph
                    textbox_blocks = self._extract_textboxes_from_element(child, block_index, table_count)
                    for tb_block in textbox_blocks:
                        self.blocks.append(tb_block)
                        block_index += 1

                para_count += 1

            elif tag == "tbl":
                # Table
                if table_count < len(self._doc.tables):
                    table = self._doc.tables[table_count]
                    table_blocks = self._extract_table(table, block_index, table_count)
                    for tb in table_blocks:
                        self.blocks.append(tb)
                        block_index += 1
                table_count += 1

        # Optionally extract headers/footers
        if self.translate_headers:
            for section_idx, section in enumerate(self._doc.sections):
                containers = (
                    ("header", section.header, "header"),
                    ("footer", section.footer, "footer"),
                    ("first_page_header", section.first_page_header, "header"),
                    ("first_page_footer", section.first_page_footer, "footer"),
                    ("even_page_header", section.even_page_header, "header"),
                    ("even_page_footer", section.even_page_footer, "footer"),
                )
                for container_name, container, block_type in containers:
                    if not container:
                        continue
                    for para_idx, para in enumerate(container.paragraphs):
                        block = self._extract_paragraph(
                            para, block_index,
                            f"section[{section_idx}].{container_name}.para[{para_idx}]",
                            block_type=block_type,
                        )
                        if block:
                            self.blocks.append(block)
                            block_index += 1

        return self.blocks

    def get_translatable_blocks(self) -> list[DocxBlock]:
        """Return only blocks that need translation."""
        return [b for b in self.blocks if b.translatable]

    def get_context_text(self, block_index: int) -> str:
        """Get previous block text for context window."""
        for i in range(block_index - 1, -1, -1):
            if self.blocks[i].translatable and self.blocks[i].text.strip():
                return self.blocks[i].text[:500]
        return ""

    # ----------------------------------------------------------
    # Paragraph extraction
    # ----------------------------------------------------------

    def _extract_paragraph(self, para, block_index: int, path: str,
                           block_type: str = "paragraph") -> Optional[DocxBlock]:
        """Extract a single paragraph with run metadata."""
        text = para.text.strip()
        if not text:
            return None

        runs_meta = []
        for run in para.runs:
            if not run.text:
                continue
            rm = RunMeta(
                text=run.text,
                bold=run.bold,
                italic=run.italic,
                underline=run.underline is not None and run.underline != False,
                font_name=run.font.name if run.font else None,
                font_size=run.font.size.pt if run.font and run.font.size else None,
                color=self._get_run_color(run),
                strike=run.font.strike if run.font else None,
            )
            runs_meta.append(rm)

        # Determine if translatable
        translatable = bool(text.strip())

        return DocxBlock(
            index=block_index,
            block_type=block_type,
            text=text,
            translatable=translatable,
            style_name=para.style.name if para.style else "",
            runs=runs_meta,
            parent_path=path,
        )

    # ----------------------------------------------------------
    # Table extraction
    # ----------------------------------------------------------

    def _extract_table(self, table, start_index: int, table_idx: int) -> list[DocxBlock]:
        """Extract all cells from a table as individual blocks."""
        blocks = []
        block_index = start_index
        seen_cells = set()

        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_key = cell._tc
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                for para_idx, para in enumerate(cell.paragraphs):
                    text = para.text.strip()
                    if not text:
                        continue

                    runs_meta = []
                    for run in para.runs:
                        if not run.text:
                            continue
                        rm = RunMeta(
                            text=run.text,
                            bold=run.bold,
                            italic=run.italic,
                            underline=run.underline is not None and run.underline != False,
                            font_name=run.font.name if run.font else None,
                            font_size=run.font.size.pt if run.font and run.font.size else None,
                            color=self._get_run_color(run),
                            strike=run.font.strike if run.font else None,
                        )
                        runs_meta.append(rm)

                    path = f"table[{table_idx}].row[{row_idx}].cell[{cell_idx}].para[{para_idx}]"
                    blocks.append(DocxBlock(
                        index=block_index,
                        block_type="table_cell",
                        text=text,
                        translatable=not _is_table_roll_marker(text),
                        style_name=para.style.name if para.style else "",
                        runs=runs_meta,
                        parent_path=path,
                        table_index=table_idx,
                        row_index=row_idx,
                        cell_index=cell_idx,
                        para_in_cell=para_idx,
                    ))
                    block_index += 1

        return blocks

    # ----------------------------------------------------------
    # Text box extraction (via XML)
    # ----------------------------------------------------------

    def _extract_textboxes_from_element(self, element, start_index: int,
                                        textbox_counter: int) -> list[DocxBlock]:
        """Extract text from text boxes (w:txbxContent) inside a paragraph element."""
        blocks = []
        block_index = start_index
        tb_idx = textbox_counter

        # Find all txbxContent elements
        txbx_elements = element.findall('.//' + qn('w:txbxContent'))
        for txbx in txbx_elements:
            para_elements = txbx.findall(qn('w:p'))
            for para_idx, p_elem in enumerate(para_elements):
                text = self._extract_text_from_p_element(p_elem)
                if not text.strip():
                    continue

                runs_meta = self._extract_runs_from_p_element(p_elem)
                path = f"textbox[{tb_idx}].para[{para_idx}]"

                blocks.append(DocxBlock(
                    index=block_index,
                    block_type="textbox_para",
                    text=text.strip(),
                    translatable=True,
                    style_name="",
                    runs=runs_meta,
                    parent_path=path,
                    textbox_index=tb_idx,
                    para_in_textbox=para_idx,
                ))
                block_index += 1
            tb_idx += 1

        return blocks

    def _extract_text_from_p_element(self, p_elem) -> str:
        """Extract plain text from a w:p XML element."""
        texts = []
        for r_elem in p_elem.findall(qn('w:r')):
            for t_elem in r_elem.findall(qn('w:t')):
                if t_elem.text:
                    texts.append(t_elem.text)
        return "".join(texts)

    def _extract_runs_from_p_element(self, p_elem) -> list[RunMeta]:
        """Extract run metadata from a w:p XML element."""
        runs = []
        for r_elem in p_elem.findall(qn('w:r')):
            text_parts = []
            for t_elem in r_elem.findall(qn('w:t')):
                if t_elem.text:
                    text_parts.append(t_elem.text)
            text = "".join(text_parts)
            if not text:
                continue

            # Parse run properties
            rPr = r_elem.find(qn('w:rPr'))
            bold = None
            italic = None
            font_name = None
            font_size = None
            color = None

            if rPr is not None:
                bold = rPr.find(qn('w:b')) is not None
                italic = rPr.find(qn('w:i')) is not None
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    font_name = rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi'))
                sz = rPr.find(qn('w:sz'))
                if sz is not None:
                    val = sz.get(qn('w:val'))
                    if val and val.isdigit():
                        font_size = int(val) / 2  # half-points to points
                color_elem = rPr.find(qn('w:color'))
                if color_elem is not None:
                    color = color_elem.get(qn('w:val'))

            runs.append(RunMeta(
                text=text,
                bold=bold,
                italic=italic,
                font_name=font_name,
                font_size=font_size,
                color=color,
            ))
        return runs

    # ----------------------------------------------------------
    # Helpers
    # ----------------------------------------------------------

    def _get_run_color(self, run) -> Optional[str]:
        """Get hex color string from a run."""
        try:
            if run.font and run.font.color and run.font.color.rgb:
                return str(run.font.color.rgb)
        except Exception:
            pass
        return None

def serialize_runs_with_markers(runs: list[RunMeta]) -> str:
    """
    Serialize runs into a text string with inline format markers.
    Used for AI translation to preserve bold/italic boundaries.

    Example output: "<b>HUMINT or Psychotherapy 40%</b>: He's talking like he saw it."
    """
    if not runs:
        return ""

    # Check if all runs have the same formatting (no markers needed)
    has_mixed_format = False
    for run in runs[1:]:
        if run.bold != runs[0].bold or run.italic != runs[0].italic:
            has_mixed_format = True
            break

    if not has_mixed_format:
        return "".join(r.text for r in runs)

    # Build marked-up text
    parts = []
    for run in runs:
        text = run.text
        if run.bold and run.italic:
            parts.append(f"<bi>{text}</bi>")
        elif run.bold:
            parts.append(f"<b>{text}</b>")
        elif run.italic:
            parts.append(f"<i>{text}</i>")
        else:
            parts.append(text)

    return "".join(parts)


def parse_marked_translation(translated: str, original_runs: list[RunMeta]) -> list[RunMeta]:
    """
    Parse a translated string with <b>/<i>/<bi> markers back into RunMeta objects.
    Preserves the format attributes from original runs.

    If parsing fails (markers don't align), falls back to single run with
    first run's formatting.
    """
    if not translated or not original_runs:
        return [RunMeta(text=translated or "")]

    # Try to parse markers
    pattern = re.compile(r'<(b|i|bi)>(.*?)</\1>', re.DOTALL)
    parts = []
    pos = 0

    for m in pattern.finditer(translated):
        # Text before this marker
        if m.start() > pos:
            parts.append(RunMeta(text=translated[pos:m.start()]))
        # Marked text
        tag = m.group(1)
        text = m.group(2)
        rm = RunMeta(
            text=text,
            bold=tag in ("b", "bi"),
            italic=tag in ("i", "bi"),
        )
        parts.append(rm)
        pos = m.end()

    # Remaining text after last marker
    if pos < len(translated):
        parts.append(RunMeta(text=translated[pos:]))

    if not parts:
        parts = [RunMeta(text=translated)]

    # Copy font metadata from original runs
    # Strategy: apply first run's font info to all parsed runs
    if original_runs:
        base = original_runs[0]
        for part in parts:
            if part.font_name is None:
                part.font_name = base.font_name
            if part.font_size is None:
                part.font_size = base.font_size
            if part.color is None:
                part.color = base.color

    return parts

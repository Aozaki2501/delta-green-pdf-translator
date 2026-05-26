from docx import Document
import pytest

from core.docx_extractor import DocxBlock, DocxExtractor
from exporters.docx_inplace import write_docx_inplace


def test_docx_extractor_skips_duplicate_merged_table_cells(tmp_path):
    path = tmp_path / "merged.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "Merged heading"
    table.cell(1, 0).text = "Left"
    table.cell(1, 1).text = "Right"
    doc.save(path)

    blocks = DocxExtractor(str(path)).extract()
    texts = [block.text for block in blocks]

    assert texts.count("Merged heading") == 1
    assert texts == ["Merged heading", "Left", "Right"]


def test_docx_extractor_reads_and_writes_first_page_footer(tmp_path):
    path = tmp_path / "footer.docx"
    output_path = tmp_path / "footer_zh.docx"
    doc = Document()
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.first_page_footer.paragraphs[0].text = "Published by arrangement."
    doc.add_paragraph("Body text")
    doc.save(path)

    blocks = DocxExtractor(str(path), translate_headers=True).extract()
    footer_block = next(block for block in blocks if block.parent_path.endswith("first_page_footer.para[0]"))

    write_docx_inplace(
        blocks,
        {footer_block.index: "经授权出版。"},
        str(path),
        str(output_path),
    )

    translated = Document(str(output_path))
    assert translated.sections[0].first_page_footer.paragraphs[0].text == "经授权出版。"


def test_docx_extractor_does_not_translate_table_roll_markers(tmp_path):
    path = tmp_path / "roll_table.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "d10 Unsettling Hassles"
    table.cell(0, 1).text = ""
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "The intercom buzzes."
    doc.save(path)

    blocks = DocxExtractor(str(path)).extract()
    translatable_texts = [block.text for block in blocks if block.translatable]

    assert "1" not in translatable_texts
    assert "d10 Unsettling Hassles" in translatable_texts
    assert "The intercom buzzes." in translatable_texts


def test_docx_inplace_raises_when_translation_target_is_missing(tmp_path):
    path = tmp_path / "source.docx"
    output_path = tmp_path / "output.docx"
    old_output = tmp_path / "old.docx"
    doc = Document()
    doc.add_paragraph("Original text")
    doc.save(path)
    old_doc = Document()
    old_doc.add_paragraph("Old output")
    old_doc.save(old_output)
    old_output.replace(output_path)

    blocks = [
        DocxBlock(
            index=1,
            block_type="paragraph",
            text="Original text",
            translatable=True,
            parent_path="body.para[99]",
        )
    ]

    with pytest.raises(RuntimeError, match="Word 写回失败"):
        write_docx_inplace(blocks, {1: "译文"}, str(path), str(output_path))

    assert Document(str(output_path)).paragraphs[0].text == "Old output"
